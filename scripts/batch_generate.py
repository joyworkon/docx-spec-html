#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import hashlib
import html
import json
import mimetypes
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from validate_output import validate


SKILL_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_STYLE = SKILL_ROOT / "assets" / "styles.css"
DEFAULT_FONT = SKILL_ROOT / "assets" / "fonts" / "JINGDONGLangZhengTi1-Bold.woff2"
DEFAULT_H2C = SKILL_ROOT / "assets" / "vendor" / "html2canvas.min.js"
DEFAULT_EDITOR = SKILL_ROOT / "assets" / "vendor" / "html-editor.html"
GENERATOR_CSS_MARKER = "/* ===== Generic DOCX generator additions ===== */"
SKILL_RELEASE = "2026.08.25-r16"


# Editorial boilerplate removed from every source: the 【官方建议】 title marker,
# the '官方建议'诠释：… block, and the 适用类目范围：… block. Each block runs from
# its label to the end of that paragraph; documents without these markers are
# left untouched.
ADVICE_TITLE_MARK = "【官方建议】"
ADVICE_BLOCK_RE = re.compile(r"['\"“”‘’]?官方建议['\"“”‘’]?诠释\s*[：:]")
SCOPE_BLOCK_RE = re.compile(r"适用类目范围\s*[：:]")


def strip_advice_boilerplate(value: str) -> str:
    text = clean_text(value).replace(ADVICE_TITLE_MARK, "")
    for pattern in (ADVICE_BLOCK_RE, SCOPE_BLOCK_RE):
        match = pattern.search(text)
        if match:
            text = text[: match.start()]
    return clean_text(text)


@dataclass
class ParagraphBlock:
    text: str
    images: list[str]
    style: str
    list_level: int | None = None  # Word numbering level (ilvl); None = not a list item


@dataclass
class TableBlock:
    table: Table


def clean_text(value: str) -> str:
    value = re.sub(r"[\u200b-\u200f\ufeff]", "", value or "")
    return re.sub(r"\s+", " ", value).strip()


def esc(value: str) -> str:
    return html.escape(value, quote=True)


JUSTIFY_ZWSP = "\u200b"
JUSTIFY_MIN_CHARS = 10  # strictly more than 10 visible characters qualifies
# Kinsoku-style guards: never break right after an opening bracket or right
# before closing punctuation when injecting break opportunities.
JUSTIFY_OPEN = set("（【《“‘〈〔「『")
JUSTIFY_CLOSE = set("），。；：、！？”’】》〉〕」』…")


def justify_word_char(char: str) -> bool:
    """ASCII printable runs (words, numbers, codes) stay atomic."""
    return 0x21 <= ord(char) <= 0x7E


def inject_justify_zwsp(text: str) -> str:
    """Insert U+200B between characters so Blink can expand justification.

    Blink's text-align: justify only expands at break opportunities (spaces);
    mixed CJK/Latin lines without spaces stay ragged. Zero-width spaces give
    the engine expansion points without visible width on the last line.
    """
    out: list[str] = []
    for index, char in enumerate(text):
        out.append(char)
        if index + 1 >= len(text):
            continue
        nxt = text[index + 1]
        if char == JUSTIFY_ZWSP or nxt == JUSTIFY_ZWSP:
            continue
        if char in JUSTIFY_OPEN or nxt in JUSTIFY_CLOSE:
            continue
        if char == " " or nxt == " ":
            continue
        if justify_word_char(char) and justify_word_char(nxt):
            continue
        out.append(JUSTIFY_ZWSP)
    return "".join(out)


def justify_visible_len(fragment: str) -> int:
    text = html.unescape(re.sub(r"<[^>]+>", "", fragment))
    return len(re.sub(r"\s", "", text))


def justify_add_class(open_tag: str, cls: str = "justify-txt") -> str:
    match = re.search(r'class="([^"]*)"', open_tag)
    if match:
        classes = match.group(1).split()
        if cls not in classes:
            classes.append(cls)
        return open_tag[: match.start(1)] + " ".join(classes) + open_tag[match.end(1) :]
    return open_tag[:-1] + f' class="{cls}">'


def justify_long_text(fragment: str) -> str:
    """Justify body text longer than 10 visible characters.

    List items get .justify-li; .label-rest paragraphs and table cells get
    .justify-txt; qualifying nodes also receive ZWSP injection. Shorter text
    keeps its existing alignment (e.g. centred table cells, .row-head).
    """

    def process_inner(inner: str) -> str:
        parts = re.split(r"(<[^>]+>)", inner)
        return "".join(part if part.startswith("<") else inject_justify_zwsp(part) for part in parts)

    def label_repl(match: re.Match) -> str:
        inner = match.group(1)
        if justify_visible_len(inner) <= JUSTIFY_MIN_CHARS:
            return match.group(0)
        return f'<div class="label-rest justify-txt">{process_inner(inner)}</div>'

    def item_repl(match: re.Match) -> str:
        open_tag, inner = match.group(1), match.group(2)
        if justify_visible_len(inner) <= JUSTIFY_MIN_CHARS:
            return match.group(0)
        return justify_add_class(open_tag, "justify-li") + process_inner(inner) + "</li>"

    def cell_repl(match: re.Match) -> str:
        open_tag, kind, inner = match.group(1), match.group(2), match.group(3)
        if justify_visible_len(inner) <= JUSTIFY_MIN_CHARS:
            return match.group(0)
        return justify_add_class(open_tag) + process_inner(inner) + f"</{kind}>"

    fragment = re.sub(r'<div class="label-rest">(.*?)</div>', label_repl, fragment, flags=re.S)
    fragment = re.sub(r"(<li(?:\s[^>]*)?>)(.*?)</li>", item_repl, fragment, flags=re.S)
    fragment = re.sub(r"(<(td|th)(?:\s[^>]*)?>)(.*?)</\2>", cell_repl, fragment, flags=re.S)
    return fragment


def slugify(value: str) -> str:
    value = clean_text(value)
    value = re.sub(r"[\\/:*?\"<>|]+", "-", value)
    value = re.sub(r"\s+", "-", value).strip("-")
    return value[:80] or "docx"


def paragraph_list_level(paragraph: Paragraph) -> int | None:
    """Word numbering level (w:numPr/w:ilvl) for the paragraph, or None when the
    paragraph is not part of a list. This is the source's own hierarchy signal:
    None = heading/top level, 0 = first list level, 1+ = nested sub-levels."""
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        return 0
    try:
        return int(ilvl.get(qn("w:val")))
    except (TypeError, ValueError):
        return 0


def iter_blocks(doc: DocumentObject) -> list[ParagraphBlock | TableBlock]:
    blocks: list[ParagraphBlock | TableBlock] = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            text = strip_advice_boilerplate(paragraph.text)
            images = paragraph_images(doc, paragraph)
            if text or images:
                blocks.append(ParagraphBlock(
                    text=text, images=images,
                    style=paragraph.style.name if paragraph.style else "",
                    list_level=paragraph_list_level(paragraph),
                ))
        elif isinstance(child, CT_Tbl):
            blocks.append(TableBlock(Table(child, doc)))
    return blocks


def paragraph_images(doc: DocumentObject, paragraph: Paragraph) -> list[str]:
    targets: list[str] = []
    for blip in paragraph._p.xpath(".//a:blip"):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid and rid in doc.part.rels:
            targets.append(doc.part.rels[rid].target_ref)
    return targets


def image_target_to_blob(doc: DocumentObject) -> dict[str, bytes]:
    blobs: dict[str, bytes] = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            blobs[rel.target_ref] = rel.target_part.blob
    return blobs


def load_css(style_path: Path, font_path: Path | None) -> str:
    """Load canonical CSS, while accepting legacy Markdown design files."""
    source = style_path.read_text(encoding="utf-8")
    if style_path.suffix.lower() == ".md":
        if "## 十七、完整 CSS 模板" in source and "## 十八、验收与交付协议" in source:
            section = source.split("## 十七、完整 CSS 模板", 1)[1].split("## 十八、验收与交付协议", 1)[0]
        else:
            section = source
        blocks = re.findall(r"```css\n(.*?)```", section, flags=re.S)
        css = "\n\n".join(blocks)
        if not css.strip():
            raise ValueError(f"No CSS code block found in {style_path}")

        # Legacy --design files contained only the base template. Preserve the
        # old behavior by appending the generator-specific component rules from
        # the canonical stylesheet.
        canonical = DEFAULT_STYLE.read_text(encoding="utf-8")
        if GENERATOR_CSS_MARKER in canonical and GENERATOR_CSS_MARKER not in css:
            css += "\n\n" + GENERATOR_CSS_MARKER + canonical.split(GENERATOR_CSS_MARKER, 1)[1]
    else:
        css = source

    if font_path and font_path.exists():
        font_data = base64.b64encode(font_path.read_bytes()).decode("ascii")
        suffix = font_path.suffix.lower()
        mime, font_format = {
            ".woff2": ("font/woff2", "woff2"),
            ".woff": ("font/woff", "woff"),
            ".otf": ("font/otf", "opentype"),
        }.get(suffix, ("font/ttf", "truetype"))
        css = re.sub(
            r'src:\s*url\("[^"]*JINGDONGLangZhengTi1-Bold\.(?:ttf|otf|woff2?|TTF|OTF|WOFF2?)"\)\s*format\("[^"]+"\);',
            f'src: url("data:{mime};base64,{font_data}") format("{font_format}");',
            css,
            count=1,
        )
    return css

EDITABLE_RUNTIME = """
<div class="edit-toolbar" data-html-edit-toolbar data-html2canvas-ignore>
  <button type="button" data-edit-toggle>编辑文字</button>
  <button type="button" data-edit-save>下载HTML</button>
  <button type="button" class="secondary" data-edit-exit>退出编辑</button>
</div>
<script>
(() => {
  const selectors = [
    "main.poster h1",
    "main.poster h2",
    "main.poster p",
    "main.poster li",
    "main.poster th",
    "main.poster td",
    "main.poster .label-text",
    "main.poster .chapter",
    "main.poster .en-label strong",
    "main.poster .hero-mark",
    "main.poster .image-caption-line span"
  ].join(",");
  const editableNodes = () => Array.from(document.querySelectorAll(selectors));
  const setEditing = (on) => {
    document.body.classList.toggle("editing", on);
    editableNodes().forEach((node) => {
      if (on) node.setAttribute("contenteditable", "true");
      else node.removeAttribute("contenteditable");
    });
    const toggle = document.querySelector("[data-edit-toggle]");
    if (toggle) toggle.textContent = on ? "正在编辑" : "编辑文字";
  };
  document.querySelector("[data-edit-toggle]")?.addEventListener("click", () => {
    setEditing(!document.body.classList.contains("editing"));
  });
  document.querySelector("[data-edit-exit]")?.addEventListener("click", () => setEditing(false));
  document.querySelector("[data-edit-save]")?.addEventListener("click", () => {
    const clone = document.documentElement.cloneNode(true);
    clone.querySelector("body")?.classList.remove("editing");
    clone.querySelectorAll("[contenteditable]").forEach((node) => node.removeAttribute("contenteditable"));
    const html = "<!doctype html>\\n" + clone.outerHTML;
    const blob = new Blob([html], { type: "text/html;charset=utf-8" });
    const link = document.createElement("a");
    link.href = URL.createObjectURL(blob);
    link.download = (document.title || "edited-output") + "-edited.html";
    link.click();
    URL.revokeObjectURL(link.href);
  });
})();
</script>
"""


def label_line(text: str) -> str:
    """Render one red-square label without splitting punctuation inside brackets.

    Module-local numbered subtitles always receive the pink marker.  A trailing
    parenthetical note that itself contains a colon (for example an upload path)
    moves intact to the next line instead of being cut at that inner colon.
    """
    t = clean_text(text)
    if LOCAL_SUBHEAD_RE.match(t):
        aside = re.match(r"^(.*?)([（(][^（）()]*[：:][^（）()]*[）)])\s*$", t)
        if aside:
            title, rest = aside.group(1).rstrip(), aside.group(2)
            return (
                f'<div class="label-line"><span class="label-text">{esc(title)}</span></div>'
                f'<div class="label-rest">{esc(rest)}</div>'
            )
        return f'<div class="label-line"><span class="label-text">{esc(t)}</span></div>'

    depth = 0
    cut = -1
    for index, char in enumerate(t):
        if char in "（(":
            depth += 1
        elif char in "）)":
            depth = max(0, depth - 1)
        elif char in "：:" and depth == 0:
            cut = index + 1
            break
    if cut != -1:
        prefix, rest = t[:cut], t[cut:].strip()
        html = f'<div class="label-line"><span class="label-text">{esc(prefix)}</span></div>'
        if rest:
            html += f'<div class="label-rest">{esc(rest)}</div>'
        return html
    return f'<div class="label-line"><span class="label-plain">{esc(t)}</span></div>'


def source_list(items: list) -> str:
    """items may be plain strings (grey level-0) or (level, text) tuples, where
    level>=1 renders as a deeper hollow-square sub-item."""
    rows = []
    for it in items:
        level, text = it if isinstance(it, tuple) else (0, it)
        if not clean_text(text):
            continue
        cls = ' class="deep"' if level and level >= 1 else ""
        rows.append(f"<li{cls}>{emphasize_prefix(esc(text))}</li>")
    return f'<ul class="source-list">{"".join(rows)}</ul>' if rows else ""


def emphasize_prefix(text: str) -> str:
    # Numbered siblings such as 1、/2、/3、must keep one typographic weight.
    # A colon inside only one sibling (e.g. 3、时长：) is not a reason to bold it.
    if NUMBERED_ITEM_RE.match(text):
        return text
    for mark in ("：", ":"):
        if mark in text and text.index(mark) <= 18:
            idx = text.index(mark) + 1
            return f"<b>{text[:idx]}</b>{text[idx:].strip()}"
    return text


def plain_block(items: list[str]) -> str:
    paras = "".join(f"<p>{esc(item)}</p>" for item in items if clean_text(item))
    return f'<div class="text-block plain-block">{paras}</div>' if paras else ""


def text_block(label: str, items: list[str]) -> str:
    return f'<div class="text-block">{label_line(label)}{source_list(items)}</div>'


def image_tag(target: str, blobs: dict[str, bytes], alt: str) -> str:
    data = blobs[target]
    mime = mimetypes.guess_type(target)[0] or "application/octet-stream"
    uri = f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"
    return f'<img class="doc-image" src="{uri}" alt="{esc(alt)}">'


def image_frame(label: str, target: str, blobs: dict[str, bytes]) -> str:
    safe_label = label or "示例图："
    return (
        '<div class="image-frame wide-image">'
        f"{label_line(safe_label)}"
        f'<div class="image-holder">{image_tag(target, blobs, safe_label)}</div>'
        "</div>"
    )


def render_table(table: Table, doc: DocumentObject, blobs: dict[str, bytes], extra_class: str = "") -> str:
    """Render a Word table without flattening merged-cell geometry.

    python-docx exposes every grid slot covered by a horizontal or vertical merge
    as a Cell wrapper around the same underlying ``w:tc`` node.  Rendering every
    wrapper duplicates both text and images.  Keep one semantic cell and express
    its grid coverage with HTML colspan/rowspan instead.
    """
    grid = [list(row.cells) for row in table.rows]
    rows_html: list[str] = []
    emitted_cells: set[int] = set()
    for row_idx, cells in enumerate(grid):
        cells_html: list[str] = []
        tag = "th" if row_idx == 0 else "td"
        cell_index = 0
        while cell_index < len(cells):
            cell = cells[cell_index]
            cell_key = id(cell._tc)
            colspan_count = 1
            while (
                cell_index + colspan_count < len(cells)
                and cells[cell_index + colspan_count]._tc is cell._tc
            ):
                colspan_count += 1
            if cell_key in emitted_cells:
                cell_index += colspan_count
                continue
            emitted_cells.add(cell_key)

            rowspan_count = 1
            while row_idx + rowspan_count < len(grid):
                below = grid[row_idx + rowspan_count]
                if cell_index >= len(below) or below[cell_index]._tc is not cell._tc:
                    break
                rowspan_count += 1

            texts = [strip_advice_boilerplate(p.text) for p in cell.paragraphs if strip_advice_boilerplate(p.text)]
            image_html = []
            for paragraph in cell.paragraphs:
                for target in paragraph_images(doc, paragraph):
                    if target in blobs:
                        image_html.append(f'<div class="image-holder">{image_tag(target, blobs, texts[0] if texts else "表格图片")}</div>')
            text_html = "<br>".join(esc(text) for text in texts)
            colspan = f' colspan="{colspan_count}"' if colspan_count > 1 else ""
            rowspan = f' rowspan="{rowspan_count}"' if rowspan_count > 1 else ""
            row_head = (
                row_idx > 0
                and cell_index == 0
                and not image_html
                and is_short_row_header(" ".join(texts))
            )
            cell_classes: list[str] = []
            if row_head:
                cell_classes.append("row-head")
            if image_html:
                cell_classes.append("table-media-cell")
            class_attr = f' class="{" ".join(cell_classes)}"' if cell_classes else ""
            cells_html.append(f"<{tag}{class_attr}{colspan}{rowspan}>{text_html}{''.join(image_html)}</{tag}>")
            cell_index += colspan_count
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    classes = "doc-table" + (f" {extra_class}" if extra_class else "")
    colgroup = ""
    if "spec-table" in extra_class and grid and len(grid[0]) == 4:
        colgroup = (
            '<colgroup><col class="spec-c1"><col class="spec-c2">'
            '<col class="spec-c3"><col class="spec-c4"></colgroup>'
        )
    return f'<div class="doc-table-wrap"><table class="{classes}">{colgroup}' + "".join(rows_html) + "</table></div>"


def document_uses_heading_styles(blocks: list[ParagraphBlock | TableBlock]) -> bool:
    """True when the DOCX marks its chapters with real Word heading styles."""
    for block in blocks:
        if isinstance(block, ParagraphBlock) and block.text:
            style = block.style.lower()
            if style.startswith("heading") or "标题" in style:
                return True
    return False


CHAPTER_LABEL_SUFFIXES = ("展示", "示例", "布局", "流程", "说明", "组合")

# 身体护理规范使用这组固定模块名；其他行业规范可有不同数量和名称。
# 对后者，脚本会把概述中的【模块名】与正文 `N、模块名` 交叉验证，避免把普通
# 编号规则误升为章节。
SPEC_MODULE_NAMES = (
    "主图规范", "主图视频", "长标题", "短标题",
    "通用卖点", "主推标签", "品质标签", "属性",
)
# 剥掉可能的 `N、`/`（一）`/`一、` 前缀后，与白名单精确比对用的正则
MODULE_NUM_PREFIX_RE = re.compile(r"^\s*[（(]?[0-9一二三四五六七八九十]+[）)、.．]\s*")


def spec_module_core(text: str) -> str:
    """剥掉一级编号前缀并去掉 （…X%）指标后，返回可与 SPEC_MODULE_NAMES 比对的核心词。"""
    core = MODULE_NUM_PREFIX_RE.sub("", text.strip())
    core = TITLE_METRIC_RE.sub("", core).strip()
    return core.rstrip("：:").strip()


def equivalent_module_name(left: str, right: str) -> bool:
    """Match overview labels to numbered chapter names without broad fuzziness."""
    aliases = {"卖点": "通用卖点", "主图规范": "主图"}
    a = aliases.get(clean_text(left), clean_text(left))
    b = aliases.get(clean_text(right), clean_text(right))
    if a == b:
        return True
    # 行业文档的缩写/扩写标题：短标↔短标题（互为前缀），
    # SPU绑定↔自营SPU销售属性绑定（概述名是章节名的子序列）。
    if len(a) >= 2 and len(b) >= 2 and (a.startswith(b) or b.startswith(a)):
        return True
    if len(a) >= 3 and len(b) >= 3:
        # 概述名是章节名的有序子序列即可（自营SPU销售属性绑定 ⊇ SPU绑定）。
        # 至少 3 个字，避免「首张主图模块化布局图」这类模块内标题误中「主图」。
        shorter, longer = (a, b) if len(a) <= len(b) else (b, a)
        it = iter(longer)
        return all(ch in it for ch in shorter)
    return False


def discover_numbered_module_headings(
    blocks: list[ParagraphBlock | TableBlock],
) -> set[str]:
    """Find document-specific numbered modules from its own overview.

    A plain Arabic-numbered paragraph is accepted only when its short title also
    appears as an overview bracket label such as ``【系列品】``.  This preserves
    industry-specific 9+ module documents without reviving the old bug where
    numbered body rules became extra chapters.

    编号也可能只存在于 Word 自动编号（w:numPr）里而文本没有 `N、` 前缀——
    例如童车规范的 12 个模块标题全部是 numId 编号、文本仅两个字到十个字。
    因此同样接受 ``list_level == 0`` 且无句子标点的短标题，前提是它仍与某个
    概述【模块名】等价（见 equivalent_module_name）。
    """
    overview_names: set[str] = set()
    for block in blocks:
        if not isinstance(block, ParagraphBlock) or not block.text:
            continue
        match = re.match(r"^\s*【([^】]+)】\s*[：:]?", block.text)
        if match:
            overview_names.add(clean_text(match.group(1)))

    discovered: set[str] = set()
    for block in blocks:
        if not isinstance(block, ParagraphBlock) or block.images or not block.text:
            continue
        match = re.match(r"^\s*\d{1,2}[、.．]\s*([^：:，。；！？,;!?]{1,20})\s*$", block.text)
        if match:
            core = clean_text(match.group(1))
        elif block.list_level == 0 and re.match(r"^[^：:，。；！？,;!?、]{1,20}$", block.text.strip()):
            # Word 自动编号短标题：编号在 numPr 中，文本本身即模块名。
            core = clean_text(block.text.strip())
        else:
            continue
        if core in SPEC_MODULE_NAMES or any(equivalent_module_name(core, name) for name in overview_names):
            discovered.add(block.text)
    return discovered
TITLE_PAREN_RE = re.compile(r"[（(][^（）()]*[)）]\s*$")
TITLE_METRIC_RE = re.compile(
    r"[（(]\s*([^（）()]*?\d+(?:\.\d+)?\s*(?:%|％|PP))\s*[)）]\s*$",
    re.I,
)


def section_core(text: str) -> str:
    """Heading text with a trailing （…）parenthetical removed, for detection."""
    return TITLE_PAREN_RE.sub("", text).strip()


def section_title_metric(text: str) -> str | None:
    """A 「（…转化率提升+X%/PP）」title suffix → its metric string for a green bar.
    A missing plus sign after 提升/增长 is normalised in so it renders green."""
    m = TITLE_METRIC_RE.search(text)
    if not m:
        return None
    return clean_text(m.group(1))  # keep source text verbatim (no injected +)


def is_section_heading(
    block: ParagraphBlock,
    heading_styles_present: bool = False,
    numbered_module_headings: set[str] | None = None,
) -> bool:
    text = block.text
    if not text or block.images:
        return False
    style = block.style.lower()
    if style.startswith("heading") or "标题" in style:
        return True
    if len(text) > 48:
        return False
    # A paragraph ending in a colon is a label/lead-in, never a chapter title.
    if text.endswith(("：", ":")):
        return False
    # 固定 8 模块白名单：剥掉 `N、`/`（一）` 前缀后精确命中白名单 → 一定是一级模块，
    # 即使带阿拉伯数字前缀（1、主图规范）也认，且优先于下面所有启发式。
    module_core = spec_module_core(text)
    if module_core in SPEC_MODULE_NAMES:
        return True
    if numbered_module_headings and text in numbered_module_headings:
        return True
    # 「整体规范综述」这类概述卡保留识别。
    if section_core(text).startswith("整体规范"):
        return True
    # Explicit chapter numerals always win, including bracketed forms like （一）.
    if re.match(r"^[（(]?[一二三四五六七八九十]+[）)、.．]", text):
        return True
    if re.match(r"^第[一二三四五六七八九十0-9]+[章节部分]", text):
        return True
    if re.match(r"^【第[一二三四五六七八九十0-9-]+屏】", text):
        return True
    # 关键根治：阿拉伯数字 `N、xxx` 编号段落**不再**无条件升级为章节。
    # 商品信息运营规范里 `1、视频画面需清晰…`、`2、视频需搭配字幕…` 都是正文
    # 编号点，只有命中上面 8 模块白名单的 `N、模块名` 才是章节。此处直接不认，
    # 无编号短语（如 `卖点选词优先级`）也因不在白名单而留作模块内子项。
    return False


def clean_section_title(text: str) -> str:
    """Strip leading chapter numerals and trailing colons from a card title."""
    cleaned = text.strip()
    cleaned = re.sub(r"^[（(][一二三四五六七八九十0-9]+[）)]\s*", "", cleaned)
    cleaned = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", cleaned)
    # 也剥掉阿拉伯数字一级编号前缀（1、主图规范 → 主图规范），外层已有 01–08 章节号
    cleaned = re.sub(r"^[0-9]+[、.．]\s*", "", cleaned)
    cleaned = re.sub(r"^第[一二三四五六七八九十0-9]+[章节部分][:：、.．\s]*", "", cleaned)
    cleaned = TITLE_METRIC_RE.sub("", cleaned).strip()  # drop only a （…X%）metric, keep asides like （建议方向）
    cleaned = cleaned.rstrip("：:").strip()
    return cleaned or text.strip()


def strip_title_prefix(title: str) -> str:
    """Remove a leading label like 主标题：/ 标题: from the document main title."""
    return re.sub(r"^\s*(?:主标题|标题|文档标题|page\s*title)\s*[:：]\s*", "", title, flags=re.I)


SCREEN_LABEL_RE = re.compile(r"^第\s*[0-9一二三四五六七八九十]+(?:\s*[-–~至]\s*[0-9一二三四五六七八九十]+)?\s*屏")


def is_label(text: str) -> bool:
    if not text or len(text) > 40:
        return False
    # A standalone "XXX：" label, allowing a trailing aside like 卖点示例组合：（无优先级）.
    if re.sub(r"[（(][^（）()]*[)）]\s*$", "", text).strip().endswith(("：", ":")):
        return True
    if re.match(r"^[0-9]{1,2}[.、．]", text):
        # A numbered line is a label only when it's a title (no sentence
        # punctuation). A numbered sentence ("1、商品品牌…一致。") is a list item.
        body = re.sub(r"^[0-9]{1,2}[.、．]\s*", "", text)
        return not re.search(r"[，。；！？,;]", body)
    if SCREEN_LABEL_RE.match(text):
        return True
    return False


def is_bare_heading(text: str) -> bool:
    """A short standalone heading with no colon and no sentence punctuation, e.g.
    補充規則 — it opens a red-square container for the deeper lines beneath it."""
    t = clean_text(text)
    return bool(t) and len(t) <= 10 and not re.search(r"[：:，。；！？,.!?]", t)


def looks_like_image_title(text: str) -> bool:
    """A short noun-phrase line that titles the picture(s) right below it, e.g.
    搜索列表页展示 / 参数楼层商详页展示 / 商品参数展示 / 示例图。Used to lift such a line
    into the image's red-square label instead of leaving it as loose body text."""
    t = clean_text(text)
    if not t or len(t) > 16 or re.search(r"[，。！？,!?]", t):
        return False
    return t.endswith(("展示", "示例", "示范", "图", "流程", "说明")) or t.endswith(("：", ":")) or len(t) <= 8


def split_sections(blocks: list[ParagraphBlock | TableBlock]) -> tuple[str, list[tuple[str, list[ParagraphBlock | TableBlock]]]]:
    first_text_index = next((i for i, block in enumerate(blocks) if isinstance(block, ParagraphBlock) and block.text), None)
    if first_text_index is None:
        return "未命名规范", [("整体规范综述", blocks, None)]
    title = strip_title_prefix(blocks[first_text_index].text)  # type: ignore[union-attr]
    body = blocks[first_text_index + 1 :]
    # Some source documents contain both a real first title paragraph and a
    # second authoring instruction such as ``标题：潮玩IP行业商品信息运营规范``.
    # The latter only tells us what belongs in the Hero; never repeat it as the
    # overview lead. Restrict this rewrite to the first body block so genuine
    # later prose beginning with “标题：” is preserved.
    if body and isinstance(body[0], ParagraphBlock):
        title_instruction = re.match(
            r"^\s*(?:主标题|标题|文档标题|page\s*title)\s*[:：]\s*(.+)$",
            body[0].text,
            flags=re.I,
        )
        if title_instruction and clean_text(title_instruction.group(1)) == clean_text(title):
            body = body[1:]
    heading_styles_present = document_uses_heading_styles(blocks)
    numbered_module_headings = discover_numbered_module_headings(blocks)
    sections: list[tuple[str, list[ParagraphBlock | TableBlock], str | None]] = []
    current_title = "整体规范综述"
    current_metric: str | None = None
    current_blocks: list[ParagraphBlock | TableBlock] = []

    for block in body:
        if isinstance(block, ParagraphBlock) and is_section_heading(
            block, heading_styles_present, numbered_module_headings
        ):
            if current_blocks:
                sections.append((current_title, current_blocks, current_metric))
            current_title = clean_section_title(block.text)
            current_metric = section_title_metric(block.text)
            current_blocks = []
        elif (
            isinstance(block, ParagraphBlock)
            and block.text
            and not block.images
            and block.text.rstrip("：:").strip() == current_title.rstrip("：:").strip()
        ):
            # Drop a paragraph that merely restates the current card title
            # (e.g. an "整体规范综述：" lead-in under the 整体规范综述 card).
            continue
        else:
            current_blocks.append(block)
    if current_blocks:
        sections.append((current_title, current_blocks, current_metric))
    if not sections:
        sections.append(("整体规范综述", body, None))
    return title, sections


BRACKET_RE = re.compile(r"^\s*(【[^】]+】)\s*(.*)$")
CIRCLED_RE = re.compile(r"^\s*[①②③④⑤⑥⑦⑧⑨⑩⑪⑫]")
LOCAL_SUBHEAD_RE = re.compile(r"^\s*[（(][0-9一二三四五六七八九十]+[）)]")
NUMBERED_ITEM_RE = re.compile(r"^\s*(?:\d+|[一二三四五六七八九十]+)[、.．]")
METRIC_UNIT_RE = r"(?:%|％|PP)"
METRIC_PAIR_RE = re.compile(
    rf"([一-龥A-Za-z·]{{2,12}})\s*([+＋]\s*\d+(?:\.\d+)?\s*{METRIC_UNIT_RE})",
    re.I,
)
# "前缀：内容" lead-in, e.g. 总结：…/字数范围：…/卖点建议顺序：… — gets a red square + pink highlight.
COLON_PREFIX_RE = re.compile(r"^([^：:\n]{1,18}[：:])(.+)$")
# A conversion metric embedded inside a longer label, e.g. "2.优化前后图 商详转化率+2%".
METRIC_INLINE_RE = re.compile(
    rf"[一-龥A-Za-z·]{{2,12}}\s*[+＋]\s*\d+(?:\.\d+)?\s*{METRIC_UNIT_RE}",
    re.I,
)


def split_label_metric(label: str) -> tuple[str | None, str | None]:
    """Pull an embedded "XX率+X%/PP" metric out of a label so it can be rendered as a
    standalone green emphasis bar, leaving the rest of the label as the title."""
    match = METRIC_INLINE_RE.search(label)
    if not match:
        return label, None
    metric = clean_text(match.group(0))
    rest = clean_text(label[: match.start()] + " " + label[match.end():]).strip(" 　：:")
    return (rest or None), metric


def is_conversion_metric(text: str) -> bool:
    """A standalone one-or-many metric line that deserves green emphasis.

    Accept either bare metrics or a source/cohort/date label followed by a colon,
    e.g. ``10SKU（0522-0531）：商详转化率+13.63%，提袋率+0.97%``.
    ``%``/``％`` and ``PP`` are equivalent metric units for component routing.
    """
    cleaned = clean_text(text)
    if not METRIC_PAIR_RE.search(cleaned):
        return False
    remainder = METRIC_PAIR_RE.sub("", cleaned)
    remainder = re.sub(r"[；;、，,\s]", "", remainder)
    if not remainder or remainder in {"效果数据", "数据效果"}:
        return True
    # Preserve the source label verbatim inside the green component; require a
    # terminal colon so ordinary prose containing a percentage is not promoted.
    return bool(re.fullmatch(r"[^：:]{1,40}[：:]", remainder))


METRIC_ARROW_SVG = (
    '<svg class="metric-arrow" xmlns="http://www.w3.org/2000/svg" '
    'viewBox="0 0 32.34917 40.82425" fill="none" aria-hidden="true">'
    '<path d="M16.405537,0L1.1559057,14.706532L11.676984,15.342317Q11.316808,35.421619,0,40.824245'
    'Q19.693825,39.944675,21.828087,15.110984L32.349167,15.74677L16.405537,0Z" fill="#47B250"/></svg>'
)

# The connected line+hook arrows rendered as INLINE svg (not CSS background) so
# html2canvas keeps them in the downloaded PNG. en-label = red, pointing left
# (flipped); hero-rule = white, pointing right.
_ARROW_PATH = (
    "M130.44582 9.5L0 9.5L0 12.5L138 12.5L138.20905 9.5146379Q137.62546 9.4325085 136.69609 9.1682091"
    "Q134.81464 8.6331568 133.25896 7.7437816Q128.52896 5.0396996 128.52898 -2.420493e-06"
    "L125.52898 2.420493e-06Q125.52896 3.7219667 127.51295 6.5661678Q128.67047 8.2256107 130.44582 9.5Z"
)
EN_LABEL_ARROW_SVG = (
    '<svg viewBox="0 0 138.209 12.5" preserveAspectRatio="none" xmlns="http://www.w3.org/2000/svg" '
    'aria-hidden="true"><g transform="translate(138.209,0) scale(-1,1)">'
    f'<path d="{_ARROW_PATH}" fill="#ff2b22"/></g></svg>'
)
HERO_RULE_ARROW_SVG = (
    '<svg viewBox="0 0 138.209 12.5" preserveAspectRatio="none" xmlns="http://www.w3.org/2000/svg" '
    f'aria-hidden="true"><path d="{_ARROW_PATH}" fill="#ffffff"/></svg>'
)

# Fully transparent 1x1 PNG: the default hero-overlay so the delivered page shows
# clean red. Replace this image in the editor (双击 → 替换图片) to float any picture
# ON TOP of the red without touching the red background colour underneath.
HERO_OVERLAY_SRC = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
)


def metric_emphasis(text: str) -> str:
    cleaned = clean_text(text)
    matches = list(METRIC_PAIR_RE.finditer(cleaned))
    if not matches:
        return f'<div class="metric-emphasis metric-standalone"><span class="metric-text">{esc(cleaned)}</span></div>'

    parts: list[str] = []
    prefix = cleaned[: matches[0].start()]
    if prefix:
        parts.append(f'<span class="metric-source-label">{esc(prefix)}</span>')
    previous_end = matches[0].start()
    for index, match in enumerate(matches):
        if index:
            separator = cleaned[previous_end : match.start()]
            if separator:
                parts.append(f'<span class="metric-separator">{esc(separator)}</span>')
        parts.append(
            '<span class="metric-item">'
            f'<span class="metric-text">{esc(match.group(1))}</span>'
            f'<span class="metric-value">{esc(match.group(2))}{METRIC_ARROW_SVG}</span>'
            '</span>'
        )
        previous_end = match.end()
    suffix = cleaned[previous_end:]
    if suffix:
        parts.append(f'<span class="metric-separator">{esc(suffix)}</span>')
    return f'<div class="metric-emphasis metric-standalone">{"".join(parts)}</div>'


# Strong "this is a clickable video" signals — trigger the 点击播放 card anywhere.
VIDEO_DEMO_RE = re.compile(r"主图视频示范|视频示范|示范视频|视频演示|视频播放按钮|播放按钮")
# Weak signals — a link or a "click to view" hint. These are only treated as a
# video play card INSIDE a 主图视频 section, so they don't hijack other sections.
VIDEO_HINT_RE = re.compile(r"查看链接|视频链接|视频地址|点击查看|点击播放|观看视频|视频入口|扫码观看|视频如下")
# A video link (http/https) provided under the 主图视频 section: its title + link
# are replaced by the 点击播放 card.
VIDEO_URL_RE = re.compile(r"https?://\S+")


def video_demo_box() -> str:
    """A 主图视频 play card: bold dark-grey 「点击播放」 label + a play icon (a solid
    dark-grey circle with a knocked-out / hollow triangle), centred on a light-pink
    ground. Triggered by a video-play placeholder line OR a video URL under the
    主图视频 card."""
    icon = (
        '<svg class="vd-icon" viewBox="0 0 48 48" xmlns="http://www.w3.org/2000/svg" aria-hidden="true">'
        # evenodd: full disc minus a triangle subpath -> the triangle is a hole.
        '<path fill-rule="evenodd" fill="#555" d="M24 1A23 23 0 1 0 24 47A23 23 0 1 0 24 1Z'
        'M19 14.5L35 24L19 33.5Z"/></svg>'
    )
    return f'<div class="video-demo"><span class="vd-text">点击播放</span>{icon}</div>'


def lead_block(text: str) -> str:
    return f'<p class="lead">{esc(text)}</p>'


def caption_line(text: str) -> str:
    """Grey-square caption (no red square, no highlight) for example images."""
    return f'<div class="caption-line">{esc(text)}</div>'


def red_list_block(items: list[str]) -> str:
    rows = []
    bracket_parent_open = False
    for item in items:
        item = clean_text(item)
        if not item:
            continue
        bracket = BRACKET_RE.match(item)
        if bracket:
            rows.append(f"<li><b>{esc(bracket.group(1))}</b>{esc(bracket.group(2))}</li>")
            bracket_parent_open = True
            continue
        class_attr = ' class="sublevel"' if bracket_parent_open else ""
        colon = COLON_PREFIX_RE.match(item)
        if colon:
            rows.append(f"<li{class_attr}><b>{esc(colon.group(1))}</b>{esc(colon.group(2).strip())}</li>")
            continue
        rows.append(f"<li{class_attr}>{esc(item)}</li>")
    if not rows:
        return ""
    return f'<div class="text-block"><ul class="red-list">{"".join(rows)}</ul></div>'


def is_short_row_header(text: str) -> bool:
    """Short first-column body labels act as vertical row headers by default."""
    compact = re.sub(r"\s+", "", clean_text(text))
    return bool(compact) and len(compact) < 10


def module_layout(items: list[str], fallback: str) -> str:
    """No-vision FALLBACK redraw of a「首张主图模块化布局图」schematic.

    This runs only when no model vision is available. It deliberately does NOT
    invent a layout (no "店铺名称" top row, no default "商品主图" / "质保承诺"
    slots — those guesses misled real models into copying a template instead of
    the actual image). It just stacks the REAL captured 主图首张 module names in a
    single column, verbatim, each in a distinct colour. An agent WITH image-reading
    ability must replace this with a faithful redraw that mirrors the real image's
    rows/columns and proportions (see references/components.md).

    Falls back to the raw image when fewer than two real module names are known."""
    names = [re.split(r"[：:]", it, 1)[0].strip() for it in items]
    names = [n for n in names if n]
    if len(names) < 2:
        return fallback

    palette = ["ml-c1", "ml-c2", "ml-c3", "ml-c4", "ml-c5", "ml-c6", "ml-c7", "ml-c8"]
    blocks = "".join(
        f'<div class="ml-block {palette[i % len(palette)]}">{esc(n)}</div>'
        for i, n in enumerate(names)
    )
    # Single-column stack of the real names — honest about what is known, with no
    # fabricated structure. The vision-capable redraw supplies the true grid.
    return (
        '<div class="module-layout"><div class="ml-grid" '
        'style="grid-template-columns:1fr;">' + blocks + "</div></div>"
    )


GENERIC_EXAMPLE_RE = re.compile(r"^示例图?\s*$")


def is_generic_example_label(label: str | None) -> bool:
    """A bare 示例 / 示例图 with NO colon and no title of its own → grey caption.
    A colon form (示例图：/ 示例：) is a real label-line and stays red."""
    return bool(label) and bool(GENERIC_EXAMPLE_RE.match(label.strip()))


def grouped_text_block(label: str | None, items: list[str], images: list[str], blobs: dict[str, bytes], half: bool, metric: str | None = None) -> str:
    """One white module that merges a label with its sub-items and example images.
    An embedded "XX率+X%" metric (pulled out of the label) renders as a green bar
    directly under the title, so it spans the module's inner width.

    Image caption rules: never synthesise a "示例图：" title. Use the source's own
    label as the picture's title — a real title (搜索列表页展示, 示例图：…) renders as a
    red-square label-line; a bare 示例/示例图 degrades to a grey caption-line."""
    generic = is_generic_example_label(label)
    inner = ""
    if label and generic:
        inner += caption_line(clean_text(label))
    elif label:
        inner += label_line(label)
    if metric:
        inner += metric_emphasis(metric)
    inner += source_list(items)
    real_images = [t for t in images if t in blobs]
    if real_images:
        # Indent the pictures one level (28px) only under a real red-square label,
        # so they align with .source-list sub-items; grey captions don't indent.
        indent = " indent" if (label and not generic) else ""
        if half:
            holder_class = "image-holder half-image" + indent
        elif len(real_images) >= 2:
            # Multiple stacked reference images: half container width, all equal.
            holder_class = "image-holder sample-image" + indent
        else:
            holder_class = "image-holder" + indent
        alt = clean_text(label) if label else "示意图"
        for target in real_images:
            inner += f'<div class="{holder_class}">{image_tag(target, blobs, alt)}</div>'
    return f'<div class="text-block">{inner}</div>' if inner else ""


def classify_table(table: Table, label: str | None = None, video_section: bool = False) -> str:
    if not table.rows:
        return "generic"
    header_cells = [clean_text(cell.text) for cell in table.rows[0].cells]
    header = " ".join(header_cells)
    all_text = " ".join(clean_text(cell.text) for row in table.rows for cell in row.cells)
    ncol = len(table.rows[0].cells)
    if video_section and VIDEO_HINT_RE.search(all_text):
        return "video_case"
    if all(name in header_cells for name in ("优化内容", "案例", "优化前", "优化后")):
        return "compare_matrix"
    if "素材图类型" in header_cells and "内容要求" in header_cells:
        return "material"
    attr_labels = {"适用肤质", "香型", "功效", "成分", "适用人群", "净含量"}
    if ncol == 2 and attr_labels.issubset(set(all_text.split())):
        return "attr"
    if "品质标签示例" in header_cells:
        return "tag_example"
    # Image-only showcase tables use one merged/equivalent heading across equal
    # columns. Route them through the shared-height media component so a narrower
    # source screenshot does not appear visibly shorter than its siblings.
    nonempty_headers = [value for value in header_cells if value]
    if (
        len(nonempty_headers) >= 2
        and len(set(nonempty_headers)) == 1
        and nonempty_headers[0] in {"展现样式", "前台展示案例"}
    ):
        return "tag_example"
    # A module-layout source is not a before/after comparison. Keep the source
    # image/text table intact for the mandatory model-led redraw step.
    if label and "首图模块化" in label:
        return "generic"
    if "优化前" in header or "优化后" in header:
        return "before_after"
    if ncol >= 3 and ("内容要求" in header or "示例" in header):
        return "spec"
    return "generic"


def before_after(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    rows = list(table.rows)
    if not rows:
        return ""
    header = [clean_text(cell.text) for cell in rows[0].cells]
    cols: list[str] = []
    for ci in range(len(rows[0].cells)):
        head = header[ci] if ci < len(header) else ""
        is_before = "前" in head or (ci == 0 and "后" not in head)
        head_class = "ba-before" if is_before else "ba-after"
        body = ""
        for row in rows[1:]:
            if ci >= len(row.cells):
                continue
            cell = row.cells[ci]
            for paragraph in cell.paragraphs:
                for target in paragraph_images(doc, paragraph):
                    if target in blobs:
                        body += f'<div class="image-holder">{image_tag(target, blobs, head)}</div>'
                txt = clean_text(paragraph.text)
                if txt:
                    body += f'<p class="ba-text">{esc(txt)}</p>'
        cols.append(f'<div class="ba-col"><div class="ba-head {head_class}">{esc(head)}</div>{body}</div>')
    return f'<div class="ba-compare">{"".join(cols)}</div>'


def spec_table(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    # Use a real table so the PDF's shared 示例 header and any other merged cells
    # remain semantic colspan/rowspan relationships.  The spec-table class only
    # selects the canonical 1fr/2fr/3fr column proportions.
    return render_table(table, doc, blobs, extra_class="spec-table")


def cell_image_targets(cell: object, doc: DocumentObject, blobs: dict[str, bytes]) -> list[str]:
    targets: list[str] = []
    for paragraph in cell.paragraphs:  # type: ignore[attr-defined]
        targets.extend(target for target in paragraph_images(doc, paragraph) if target in blobs)
    return targets


def compare_matrix(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    """Restore the PDF's grouped 4-column optimisation matrix."""
    rows = list(table.rows)
    if len(rows) < 2 or len(rows[0].cells) < 4:
        return render_table(table, doc, blobs)
    headers = [clean_text(cell.text) for cell in rows[0].cells[:4]]
    body: list[str] = []
    index = 1
    while index < len(rows):
        group_text = clean_text(rows[index].cells[0].text)
        group_end = index + 1
        while group_end < len(rows) and clean_text(rows[group_end].cells[0].text) == group_text:
            group_end += 1
        span = group_end - index
        for row_index in range(index, group_end):
            row = rows[row_index]
            cells: list[str] = []
            if row_index == index:
                cells.append(f'<td class="cm-group" rowspan="{span}">{esc(group_text)}</td>')
            case_text = clean_text(row.cells[1].text)
            cells.append(f'<td class="cm-case">{esc(case_text)}</td>')

            left_text = clean_text(row.cells[2].text)
            right_text = clean_text(row.cells[3].text)
            if case_text == "优化总结" and left_text and left_text == right_text:
                paragraphs = [strip_advice_boilerplate(p.text) for p in row.cells[2].paragraphs if strip_advice_boilerplate(p.text)]
                content = "".join(f"<p>{esc(text)}</p>" for text in paragraphs)
                cells.append(f'<td class="cm-text" colspan="2">{content}</td>')
            else:
                for column in (2, 3):
                    cell = row.cells[column]
                    targets = cell_image_targets(cell, doc, blobs)
                    texts = [strip_advice_boilerplate(p.text) for p in cell.paragraphs if strip_advice_boilerplate(p.text)]
                    classes = "cm-img" if targets else "cm-text"
                    content = "".join(
                        f'<div class="image-holder">{image_tag(target, blobs, headers[column])}</div>'
                        for target in targets
                    )
                    content += "".join(f"<p>{esc(text)}</p>" for text in texts)
                    cells.append(f'<td class="{classes}">{content}</td>')
            body.append("<tr>" + "".join(cells) + "</tr>")
        index = group_end
    colgroup = '<colgroup><col class="cm-c1"><col class="cm-c2"><col class="cm-c3"><col class="cm-c4"></colgroup>'
    head = "<thead><tr>" + "".join(f"<th>{esc(text)}</th>" for text in headers) + "</tr></thead>"
    return f'<table class="compare-matrix">{colgroup}{head}<tbody>{"".join(body)}</tbody></table>'


def material_table(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    rows = list(table.rows)
    if len(rows) < 2 or len(rows[0].cells) < 4:
        return spec_table(table, doc, blobs)
    body: list[str] = []
    for row in rows[1:]:
        kind = clean_text(row.cells[0].text)
        requirements = [clean_text(p.text) for p in row.cells[1].paragraphs if clean_text(p.text)]
        cells = [f'<td class="mt-type row-head">{esc(kind)}</td>', f'<td class="mt-req">{"<br>".join(esc(x) for x in requirements)}</td>']
        for column in (2, 3):
            targets = cell_image_targets(row.cells[column], doc, blobs)
            content = "".join(
                f'<div class="image-holder">{image_tag(target, blobs, kind + "示例")}</div>'
                for target in targets
            )
            cells.append(f'<td class="mt-eg">{content}</td>')
        body.append("<tr>" + "".join(cells) + "</tr>")
    return (
        '<table class="material-table">'
        '<colgroup><col class="mt-c1"><col class="mt-c2"><col class="mt-c3"><col class="mt-c4"></colgroup>'
        '<thead><tr><th>素材图类型</th><th>内容要求</th><th colspan="2">示例</th></tr></thead>'
        f'<tbody>{"".join(body)}</tbody></table>'
    )


def attribute_direction_table(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    rows = list(table.rows)
    if len(rows) < 6:
        return render_table(table, doc, blobs)
    body: list[str] = []
    for row_index in range(0, len(rows), 2):
        if row_index + 1 >= len(rows):
            break
        labels = [clean_text(cell.text) for cell in rows[row_index].cells[:2]]
        body.append("<tr>" + "".join(f'<th class="attr-head-red">{esc(label)}</th>' for label in labels) + "</tr>")
        image_cells: list[str] = []
        for column in (0, 1):
            targets = cell_image_targets(rows[row_index + 1].cells[column], doc, blobs)
            content = "".join(
                f'<div class="image-holder">{image_tag(target, blobs, labels[column])}</div>'
                for target in targets
            )
            image_cells.append(f'<td class="attr-img">{content}</td>')
        body.append("<tr>" + "".join(image_cells) + "</tr>")
    return f'<table class="attr-table"><tbody>{"".join(body)}</tbody></table>'


def video_case_table(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    rows = list(table.rows)
    if len(rows) < 2 or len(rows[0].cells) < 2:
        return video_demo_box()
    copy_cell, media_cell = rows[1].cells[:2]
    paragraphs = [clean_text(p.text) for p in copy_cell.paragraphs if clean_text(p.text)]
    copy_html = "".join(f'<p class="ba-text">{esc(text)}</p>' for text in paragraphs)
    targets = cell_image_targets(media_cell, doc, blobs)
    media_html = video_demo_box() + "".join(
        f'<div class="image-holder">{image_tag(target, blobs, "视频案例")}</div>' for target in targets
    )
    return (
        '<div class="video-case-grid video-case-card">'
        f'<div class="video-case-copy"><div class="video-case-head">视频案例</div><div class="video-case-body">{copy_html}</div></div>'
        f'<div class="video-case-media">{media_html}</div>'
        '</div>'
    )


def table_component(
    table: Table,
    doc: DocumentObject,
    blobs: dict[str, bytes],
    *,
    label: str | None = None,
    video_section: bool = False,
) -> tuple[str, str]:
    """Return the semantic table kind and its unwrapped component HTML."""
    kind = classify_table(table, label=label, video_section=video_section)
    if kind == "before_after":
        inner = before_after(table, doc, blobs)
    elif kind == "spec":
        inner = spec_table(table, doc, blobs)
    elif kind == "compare_matrix":
        inner = compare_matrix(table, doc, blobs)
    elif kind == "material":
        inner = material_table(table, doc, blobs)
    elif kind == "attr":
        inner = attribute_direction_table(table, doc, blobs)
    elif kind == "video_case":
        inner = video_case_table(table, doc, blobs)
    elif kind == "tag_example":
        inner = render_table(table, doc, blobs, extra_class="tag-example-table")
    else:
        inner = render_table(table, doc, blobs)
    return kind, inner


def table_group(
    label: str | None,
    table: Table,
    doc: DocumentObject,
    blobs: dict[str, bytes],
    metric: str | None = None,
    video_section: bool = False,
) -> str:
    kind, inner = table_component(
        table,
        doc,
        blobs,
        label=label,
        video_section=video_section,
    )
    if kind == "generic" and not label:
        return inner
    label_html = label_line(label) if label else ""
    # Embedded metric (e.g. 商详转化率+2%) sits under the title, spanning the same
    # width as the before/after (优化前+优化后) columns below it.
    metric_html = metric_emphasis(metric) if metric else ""
    if kind == "tag_example" and not label_html and not metric_html:
        return inner
    return f'<div class="text-block">{label_html}{metric_html}{inner}</div>'


def render_section_blocks(blocks: list[ParagraphBlock | TableBlock], doc: DocumentObject, blobs: dict[str, bytes], *, is_intro: bool = False, half_images: bool = False, video_section: bool = False) -> str:
    rendered: list[str] = []
    plain_items: list[str] = []
    bracket_items: list[str] = []
    pending_label: str | None = None
    pending_metric: str | None = None
    pending_items: list[str] = []
    pending_images: list[str] = []
    module_items: list[str] = []  # captured 主图首张 module names, for the layout redraw
    lead_done = False
    video_card_done = False  # in a 主图视频 section, emit the play card only once
    consumed_until = -1

    def flush_plain() -> None:
        nonlocal plain_items
        if plain_items:
            rendered.append(plain_block(plain_items))
            plain_items = []

    def flush_bracket() -> None:
        nonlocal bracket_items
        if bracket_items:
            rendered.append(red_list_block(bracket_items))
            bracket_items = []

    def flush_label() -> None:
        nonlocal pending_label, pending_metric, pending_items, pending_images, module_items
        if pending_label is None and not pending_images and not pending_metric and not pending_items:
            return
        item_texts = [it[1] if isinstance(it, tuple) else it for it in pending_items]
        # Capture the 主图首张 module list so a later 模块化布局图 can be redrawn.
        if pending_label and "主图首张" in pending_label:
            mods = [t for t in item_texts if ("：" in t or ":" in t)]
            if mods:
                module_items = mods
        # Redraw "首张主图模块化布局图" as a clean module schematic (no watermark)
        # instead of embedding the raw reference screenshot.
        if pending_label and "模块化布局图" in pending_label and pending_images and module_items:
            real = [t for t in pending_images if t in blobs]
            fallback = "".join(
                f'<div class="image-holder">{image_tag(t, blobs, "模块化布局图")}</div>' for t in real
            )
            ml = module_layout(module_items, fallback)
            rendered.append(f'<div class="text-block">{label_line(pending_label)}{ml}</div>')
        else:
            rendered.append(grouped_text_block(pending_label, pending_items, pending_images, blobs, half_images, pending_metric))
        pending_label = None
        pending_metric = None
        pending_items = []
        pending_images = []

    def deeper_follows(idx: int, cur_level: int | None) -> bool:
        """True when the next content line is one hierarchy step deeper than the
        current line — i.e. the current line is a heading that opens a grey
        container. A deeper Word list level, a manual ①②③, or an image all count;
        a same/shallower level or an end means the current line is a sibling item."""
        base = -1 if cur_level is None else cur_level
        for nb in blocks[idx + 1:]:
            if isinstance(nb, TableBlock):
                return False
            if not nb.text:
                if nb.images:
                    return True
                continue
            if CIRCLED_RE.match(nb.text):
                return True
            if NUMBERED_ITEM_RE.match(nb.text):
                return True
            return nb.list_level is not None and nb.list_level > base
        return False

    def table_follows(idx: int) -> bool:
        """True when the next non-empty block is a Word table — used to lift a
        caption-like line (e.g. 完整主图结构规范示例) into that table's title."""
        for nb in blocks[idx + 1:]:
            if isinstance(nb, TableBlock):
                return True
            if isinstance(nb, ParagraphBlock):
                if nb.text or nb.images:
                    return False
                continue
            return False
        return False

    def subtitle_table_pairs(idx: int) -> tuple[list[tuple[str, Table]], int]:
        """Collect consecutive ``child label → table`` pairs after a local subtitle.

        These pairs are one semantic child group in the PDF and must remain in
        the subtitle's white container. Their labels render as grey-square child
        headings rather than separate red/pink modules.
        """
        pairs: list[tuple[str, Table]] = []
        cursor = idx + 1
        while cursor + 1 < len(blocks):
            label_block = blocks[cursor]
            table_block = blocks[cursor + 1]
            if not (
                isinstance(label_block, ParagraphBlock)
                and label_block.text
                and not label_block.images
                and is_label(label_block.text)
                and isinstance(table_block, TableBlock)
            ):
                break
            pairs.append((label_block.text, table_block.table))
            cursor += 2
        return pairs, cursor

    for idx, block in enumerate(blocks):
        if idx <= consumed_until:
            continue
        if isinstance(block, TableBlock):
            flush_plain()
            flush_bracket()
            if pending_label is not None and not pending_items and not pending_images:
                label = pending_label
                metric = pending_metric
                pending_label = None
                pending_metric = None
                rendered.append(table_group(label, block.table, doc, blobs, metric=metric, video_section=video_section))
            else:
                flush_label()
                rendered.append(table_group(None, block.table, doc, blobs, video_section=video_section))
            continue

        if block.images and not block.text:
            flush_bracket()
            # If the picture is directly preceded by a short title line (展示/示例/
            # 图…), lift it out of the loose text into the image's red-square label.
            if pending_label is None and not pending_images and plain_items and looks_like_image_title(plain_items[-1]):
                pending_label = plain_items.pop()
            flush_plain()
            # Accumulate into pending_images so consecutive image-only paragraphs
            # land in ONE module and share a single uniform width class.
            pending_images.extend(block.images)
            continue

        text = block.text
        if not text:
            continue

        if is_conversion_metric(text):
            flush_plain()
            flush_bracket()
            flush_label()
            rendered.append(metric_emphasis(text))
            continue

        is_video_strong = bool(VIDEO_DEMO_RE.search(text))
        is_video_weak = video_section and not is_video_strong and bool(
            VIDEO_URL_RE.search(text) or VIDEO_HINT_RE.search(text)
        )
        if is_video_strong or is_video_weak:
            flush_plain()
            flush_bracket()
            if is_video_weak:
                # The card replaces the title + link/hint: drop a pending bare title
                # (e.g. 主图视频：/ 视频链接：/ 查看链接) instead of rendering it.
                pending_label = None
                pending_metric = None
                pending_items = []
                pending_images = []
            else:
                flush_label()
            # In a 主图视频 section collapse repeated placeholders/links to one card.
            if not (video_section and video_card_done):
                rendered.append(video_demo_box())
                video_card_done = True
            continue

        # Numbered local subtitles such as （1）…（4） always open a pink-marked
        # child module inside the current chapter. They may introduce a label
        # and table several blocks later, so requiring an immediately following
        # table incorrectly demotes the first subtitle to a grey list item.
        if LOCAL_SUBHEAD_RE.match(text):
            flush_plain()
            flush_bracket()
            flush_label()
            pairs, next_index = subtitle_table_pairs(idx)
            if pairs:
                group_inner = label_line(text)
                for child_label, child_table in pairs:
                    _, child_table_html = table_component(
                        child_table,
                        doc,
                        blobs,
                        label=child_label,
                        video_section=video_section,
                    )
                    group_inner += (
                        '<div class="nested-table-group">'
                        f'{caption_line(child_label)}{child_table_html}'
                        '</div>'
                    )
                rendered.append(
                    f'<div class="text-block subtitle-table-group">{group_inner}</div>'
                )
                consumed_until = next_index - 1
                continue
            pending_label = text
            pending_items = []
            pending_images = []
            continue

        # A caption-like line immediately followed by a TABLE is that table's
        # title: lift it to a red-square label-line in the SAME module as the
        # table (e.g. 完整主图结构规范示例 above the 主图/内容要求/示例 spec table),
        # even when a previous container is still open. Without this the line is
        # swallowed as a grey sub-item of the open container and the table loses
        # its title.
        if looks_like_image_title(text) and table_follows(idx):
            flush_plain()
            flush_bracket()
            flush_label()
            pending_label = text
            pending_items = []
            pending_images = []
            continue

        if BRACKET_RE.match(text):
            flush_plain()
            flush_label()
            bracket_items.append(text)
            continue

        # A colon label opens a new red container — UNLESS it is itself a Word list
        # item nested inside an already-open container (e.g. 性能及其他利益点： at
        # ilvl=0 under 2.详细规范：), in which case it stays a grey sub-item below.
        if is_label(text) and not (
            pending_label is not None
            and (block.list_level is not None or NUMBERED_ITEM_RE.match(text))
        ):
            flush_plain()
            flush_bracket()
            flush_label()
            rest, metric = split_label_metric(text)
            pending_label = rest
            pending_metric = metric
            pending_items = []
            pending_images = []
            continue

        # A short noun-phrase line that titles the picture(s) right below it
        # (搜索列表页展示 / 参数楼层商详页展示 / 商品参数展示) opens a fresh red-square
        # image module, breaking any container or image group currently open.
        if block.list_level is None and looks_like_image_title(text) and deeper_follows(idx, None):
            flush_plain()
            flush_bracket()
            flush_label()
            pending_label = text
            pending_items = []
            pending_images = []
            continue

        # Inside an open red container, every line is a grey sub-detail whose
        # indent comes straight from its Word level (ilvl); an ilvl=None line sits
        # at level 0.
        if pending_label is not None:
            level = block.list_level if block.list_level is not None else 0
            pending_items.append((level, text))
            continue

        # No container open: this line is top-level.
        flush_label()
        if is_intro and not lead_done and not rendered and not plain_items and not bracket_items:
            rendered.append(lead_block(text))
            lead_done = True
        elif (block.list_level is None and (is_bare_heading(text) or COLON_PREFIX_RE.match(text))
              and deeper_follows(idx, None)):
            # An ilvl=None heading (補充規則, or 前缀：… with deeper lines under it)
            # opens a red container; the deeper lines become grey sub-items.
            flush_plain()
            flush_bracket()
            pending_label = text
            pending_items = []
            pending_images = []
        elif COLON_PREFIX_RE.match(text) or block.list_level is not None:
            # A standalone "前缀：内容" line, OR any Word-bulleted line sitting
            # directly under the chapter (e.g. 参数楼层&属性 的前两段) → red-square
            # item. No colon ⇒ red square only, no pink bar.
            flush_plain()
            bracket_items.append(text)
        else:
            # A plain unbulleted paragraph with no heading role → its own white box.
            flush_bracket()
            plain_items.append(text)

    flush_plain()
    flush_bracket()
    flush_label()
    return "".join(rendered) or plain_block([""])


def hero_title_html(title: str) -> str:
    """Force the "商品信息运营规范" category suffix onto the hero title's second line."""
    title = strip_title_prefix(title)
    marker = "商品信息运营规范"
    if marker in title:
        prefix = title[: title.index(marker)].rstrip()
        return f"{esc(prefix)}<br>{esc(marker)}" if prefix else esc(marker)
    return esc(title)


def section_head(num: int | None, title: str) -> str:
    chapter = f'<span class="chapter">{num:02d}</span>' if num is not None else ""
    return (
        '<div class="section-head spec-head">'
        f"{chapter}"
        f"<h2>{{ {esc(title)} }}</h2>"
        f'<div class="en-label"><strong>INTRODUCTION</strong><span>{EN_LABEL_ARROW_SVG}</span></div>'
        "</div>"
    )


def normalize_update_label(value: str | None, docx_path: Path | None = None) -> str:
    if not value:
        # Default to the source file's last-modified date, to the day.
        d = date.fromtimestamp(docx_path.stat().st_mtime) if docx_path is not None else date.today()
        return f"更新日期 {d.year}年{d.month}月{d.day}日"
    text = clean_text(value)
    if re.search(r"[\u4e00-\u9fff]", text):
        return text
    match = re.search(r"(20\d{2})[.\-/年 ]+(\d{1,2})(?:[.\-/月 ]+(\d{1,2}))?", text)
    if match:
        year, month, day = match.group(1), int(match.group(2)), match.group(3)
        if day:
            return f"更新日期 {year}年{month}月{int(day)}日"
        return f"更新日期 {year}年{month}月"
    return f"更新日期 {text}"


def render_card(title: str, body: str, num: int | None, metric: str | None = None) -> str:
    card_class = "card intro-card" if num is None else "card spec-card"
    bar = metric_emphasis(metric) if metric else ""
    return f'<section class="{card_class}">{section_head(num, title)}<div class="gray-panel spec-text">{bar}{body}</div></section>'


# Shared helper: html2canvas does NOT support object-fit or mix-blend-mode, so a
# replaced .hero-overlay image would download deformed (stretched) and un-blended
# (wrong colour). Before capture we pre-composite the red + overlay image into a
# single flat, hero-sized PNG (cover-fit + the blend baked in via canvas
# globalCompositeOperation), swap it in (no blend, fill = no distortion), and
# restore afterwards — so the download matches the on-screen preview.
HERO_FLATTEN_JS = (
    "function flattenHeroOverlay(doc){return new Promise(function(resolve){try{"
    "var ov=doc.querySelector('.hero-overlay');"
    "if(!ov||!ov.naturalWidth||ov.naturalWidth<=2){return resolve(null);}"
    "var hero=ov.closest('.hero')||ov.parentNode;var r=hero.getBoundingClientRect();"
    "var W=Math.max(1,Math.round(r.width)),H=Math.max(1,Math.round(r.height));"
    "var win=doc.defaultView||window,cs=win.getComputedStyle(ov);"
    "var blend=(cs.mixBlendMode&&cs.mixBlendMode!=='normal')?cs.mixBlendMode:'overlay';"
    "var bg=win.getComputedStyle(hero).backgroundColor||'#FF2B22';"
    "var cnv=doc.createElement('canvas');cnv.width=W;cnv.height=H;var ctx=cnv.getContext('2d');"
    "ctx.fillStyle=bg;ctx.fillRect(0,0,W,H);"
    "var iw=ov.naturalWidth,ih=ov.naturalHeight,sc=Math.max(W/iw,H/ih),dw=iw*sc,dh=ih*sc;"
    "ctx.globalCompositeOperation=blend;ctx.drawImage(ov,(W-dw)/2,(H-dh)/2,dw,dh);"
    "ctx.globalCompositeOperation='source-over';var flat=cnv.toDataURL('image/png');"
    "var saved={src:ov.getAttribute('src'),mb:ov.style.mixBlendMode,of:ov.style.objectFit};"
    "var fired=false;var done=function(){if(fired)return;fired=true;ov.removeEventListener('load',done);"
    "resolve(function(){ov.setAttribute('src',saved.src);ov.style.mixBlendMode=saved.mb;ov.style.objectFit=saved.of;});};"
    "ov.addEventListener('load',done);ov.style.mixBlendMode='normal';ov.style.objectFit='fill';"
    "ov.setAttribute('src',flat);setTimeout(done,800);}catch(e){resolve(null);}});}"
)

# html2canvas has no <video> handling: a video only renders if the browser already
# decoded a frame. During capture, swap each <video> for an <img> of its poster
# (or decoded frame) so the video's cover always appears, then restore.
VIDEO_SWAP_JS = (
    "function swapVideosForPosters(doc){var swaps=[];"
    "[].slice.call(doc.querySelectorAll('video')).forEach(function(v){"
    "var src=v.getAttribute('poster')||'';"
    "if(!src&&v.videoWidth){try{var c=doc.createElement('canvas');c.width=v.videoWidth;c.height=v.videoHeight;"
    "c.getContext('2d').drawImage(v,0,0,c.width,c.height);src=c.toDataURL('image/jpeg',0.85);}catch(e){}}"
    "if(!src)return;var img=doc.createElement('img');img.src=src;img.className=v.className;"
    "img.style.cssText=v.style.cssText;img.style.display='block';img.style.width=v.style.width||'100%';img.style.height='auto';"
    "v.parentNode.insertBefore(img,v);var pd=v.style.display;v.style.display='none';swaps.push({v:v,img:img,pd:pd});});"
    "return function(){swaps.forEach(function(s){if(s.img.parentNode)s.img.parentNode.removeChild(s.img);s.v.style.display=s.pd;});};}"
)


def download_runtime() -> str:
    """A floating "下载整页图片" button that rasterises the whole poster to one
    PNG via an embedded html2canvas, so the page stays self-contained/offline.
    Scale is clamped so the canvas height stays under the browser limit."""
    if not DEFAULT_H2C.exists():
        return ""
    lib_b64 = base64.b64encode(DEFAULT_H2C.read_bytes()).decode("ascii")
    return (
        '<button id="dl-page-btn" class="dl-page-btn" data-html2canvas-ignore>下载整页图片</button>\n'
        f'<script type="application/octet-stream" id="html2canvas-src-b64" data-html2canvas-ignore>{lib_b64}</script>\n'
        "<script>(function(){var b=document.getElementById('dl-page-btn');"
        "var srcEl=document.getElementById('html2canvas-src-b64');"
        "function ensureH2C(){if(window.html2canvas){if(b)b.setAttribute('data-h2c-status','ready');return true;}"
        "if(b)b.setAttribute('data-h2c-status','loading');if(!srcEl){if(b)b.setAttribute('data-h2c-status','missing-source');return false;}"
        "try{var bin=atob(srcEl.textContent.trim());var bytes=new Uint8Array(bin.length);"
        "for(var i=0;i<bin.length;i++){bytes[i]=bin.charCodeAt(i);}"
        "var lib=new TextDecoder('utf-8').decode(bytes);"
        "var s=document.createElement('script');s.text='(function(){var module=undefined,exports=undefined,define=undefined;'+lib+'\\n}).call(window);';"
        "document.head.appendChild(s);s.remove();var ok=!!window.html2canvas;if(b)b.setAttribute('data-h2c-status',ok?'ready':'unregistered');return ok;}"
        "catch(e){if(b)b.setAttribute('data-h2c-status','error');console.error(e);return false;}}"
        "if(!b||!ensureH2C())return;"
        f"{HERO_FLATTEN_JS}{VIDEO_SWAP_JS}"
        "b.addEventListener('click',function(){var p=document.querySelector('.poster');if(!p)return;"
        "var t=b.textContent;b.textContent='生成中…';b.disabled=true;var restore=null,vrestore=null;"
        "flattenHeroOverlay(document).then(function(rf){restore=rf;vrestore=swapVideosForPosters(document);"
        "var h=p.scrollHeight,s=Math.min(2,Math.max(0.3,32000/h));"
        "return html2canvas(p,{scale:s,backgroundColor:'#dcedff',useCORS:true,logging:false,"
        "windowWidth:p.scrollWidth,windowHeight:h});}).then(function(c){if(vrestore)vrestore();if(restore)restore();c.toBlob(function(bl){"
        "var a=document.createElement('a');a.href=URL.createObjectURL(bl);"
        "a.download=(document.title||'page')+'.png';document.body.appendChild(a);a.click();a.remove();"
        "setTimeout(function(){URL.revokeObjectURL(a.href);},1500);b.textContent=t;b.disabled=false;},'image/png');"
        "}).catch(function(e){if(vrestore)vrestore();if(restore)restore();console.error(e);b.textContent='下载失败，重试';b.disabled=false;});});})();</script>"
    )


def editor_runtime() -> str:
    """A floating 「编辑」 button that opens the bundled visual HTML editor in a NEW
    window, pre-loaded with the current page. The whole editor is embedded as inert
    base64 so the page stays self-contained / offline (no second file needed)."""
    if not DEFAULT_EDITOR.exists():
        return ""
    editor_bytes = DEFAULT_EDITOR.read_bytes()
    editor_sha256 = hashlib.sha256(editor_bytes).hexdigest()
    b64 = base64.b64encode(editor_bytes).decode("ascii")
    launcher = (
        "(function(){var b=document.getElementById('edit-page-btn');"
        "var srcEl=document.getElementById('editor-src-b64');"
        "if(!b||!srcEl)return;"
        "b.addEventListener('click',function(){"
        # Clean snapshot of the current poster: drop the floating controls + scripts.
        "var root=document.documentElement.cloneNode(true);"
        "root.querySelectorAll('[data-html2canvas-ignore],script').forEach(function(n){n.remove();});"
        "var poster='<!doctype html>\\n'+root.outerHTML;"
        # Decode the embedded editor (UTF-8 safe), then run it in a new window with
        # the current page handed off via window.__PRELOAD_HTML__.
        "var bin=atob(srcEl.textContent.trim());var bytes=new Uint8Array(bin.length);"
        "for(var i=0;i<bin.length;i++){bytes[i]=bin.charCodeAt(i);}"
        "var editorHtml=new TextDecoder('utf-8').decode(bytes);"
        "var name=(document.title||'页面')+'.html';"
        "var boot='<script>window.__PRELOAD_NAME__='+JSON.stringify(name)+';window.__PRELOAD_HTML__='+JSON.stringify(poster)+';<\\/script>';"
        # Inject boot inside <head> so the doctype stays first (no quirks mode).
        "var out=editorHtml.replace(/<head([^>]*)>/i,function(m){return m+boot;});"
        "if(out===editorHtml){out=boot+editorHtml;}"
        "var w=window.open('','_blank');"
        "if(!w){alert('请允许弹出窗口后重试');return;}"
        "w.document.open();w.document.write(out);w.document.close();"
        "});})();"
    )
    return (
        '<button id="edit-page-btn" class="edit-page-btn" data-html2canvas-ignore>编辑</button>\n'
        f'<script type="application/octet-stream" id="editor-src-b64" data-editor-sha256="{editor_sha256}" data-html2canvas-ignore>{b64}</script>\n'
        f"<script>{launcher}</script>"
    )


def render_html(docx_path: Path, style_path: Path, font_path: Path | None, updated_label: str, editable: bool) -> str:
    doc = Document(docx_path)
    blobs = image_target_to_blob(doc)
    title, sections = split_sections(iter_blocks(doc))
    css = load_css(style_path, font_path)
    cards: list[str] = []
    chapter = 1
    for section_title, section_blocks, section_metric in sections:
        is_intro = section_title == "整体规范综述" and not cards
        half_images = "图文详情" in section_title
        video_section = "主图视频" in section_title
        body = render_section_blocks(
            section_blocks, doc, blobs, is_intro=is_intro, half_images=half_images,
            video_section=video_section,
        )
        body = justify_long_text(body)
        if is_intro:
            cards.append(render_card(section_title, body, None))
        else:
            cards.append(render_card(section_title, body, chapter, section_metric))
            chapter += 1
    editable_runtime = EDITABLE_RUNTIME if editable else ""
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <meta name="generator" content="docx-spec-html/{SKILL_RELEASE}">
  <title>{esc(title)}</title>
  <style>{css}</style>
</head>
<body>
<main class="poster auto-doc">
  <section class="hero">
    <img class="hero-overlay" src="{HERO_OVERLAY_SRC}" alt="头图背景图（在编辑器中双击此处可替换为叠加在红底上的图片）">
    <h1>{hero_title_html(title)}</h1>
    <p class="updated">{esc(updated_label)}</p>
    <div class="hero-mark">OPERATION<br>STANDARDS</div>
    <div class="hero-rule">{HERO_RULE_ARROW_SVG}</div>
  </section>
  {''.join(cards)}
</main>
{download_runtime()}
{editor_runtime()}
{editable_runtime}
</body>
</html>
"""


def docx_inputs(input_path: Path) -> list[Path]:
    if input_path.is_file():
        return [input_path]
    return sorted(path for path in input_path.rglob("*.docx") if not path.name.startswith("~$"))


def generate_one(docx_path: Path, output_dir: Path, style_path: Path, font_path: Path | None, updated_value: str | None, editable: bool) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    out_html = output_dir / f"{slugify(docx_path.stem)}-output.html"
    report_path = output_dir / f"{slugify(docx_path.stem)}-report.json"
    updated_label = normalize_update_label(updated_value, docx_path)
    html_text = render_html(docx_path, style_path, font_path, updated_label, editable)
    out_html.write_text(html_text, encoding="utf-8")
    report = validate(docx_path, out_html)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return {
        "source": str(docx_path),
        "html": str(out_html),
        "report": str(report_path),
        "passed": report["passed"],
        "warnings": report["warnings"],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Batch-generate single-file HTML from DOCX files.")
    parser.add_argument("input", type=Path, help="A .docx file or a folder containing .docx files.")
    parser.add_argument("output_dir", type=Path, help="Folder for generated HTML and reports.")
    parser.add_argument(
        "--style", "--design", dest="style", type=Path, default=DEFAULT_STYLE,
        help="CSS stylesheet path. --design remains as a compatibility alias and also accepts legacy Markdown design files.",
    )
    parser.add_argument("--font", type=Path, default=DEFAULT_FONT)
    parser.add_argument("--updated", default=None, help='Hero update label. Examples: "2026.06" or "更新日期 2026年06月".')
    parser.add_argument(
        "--editable",
        action="store_true",
        help="Add the optional inline contenteditable toolbar; omit for standard delivery, which already includes 编辑 and 下载整页图片.",
    )
    parser.add_argument("--strict", action="store_true", help="Exit non-zero if any generated report has warnings.")
    args = parser.parse_args()

    inputs = docx_inputs(args.input)
    if not inputs:
        raise SystemExit(f"No .docx files found: {args.input}")

    font_path = args.font if args.font and args.font.exists() else None
    summary = [generate_one(path, args.output_dir, args.style, font_path, args.updated, args.editable) for path in inputs]
    summary_path = args.output_dir / "batch-summary.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"count": len(summary), "summary": str(summary_path), "items": summary}, ensure_ascii=False, indent=2))
    return 1 if args.strict and any(not item["passed"] for item in summary) else 0


if __name__ == "__main__":
    raise SystemExit(main())
