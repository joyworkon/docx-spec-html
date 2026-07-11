#!/usr/bin/env python3
from __future__ import annotations

import argparse
import html.parser
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document


def clean_text(value: str) -> str:
    value = re.sub(r"[\u200b-\u200f\ufeff]", "", value or "")
    return re.sub(r"\s+", " ", value).strip()


class VisibleTextParser(html.parser.HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"style", "script", "noscript"}:
            self.skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"style", "script", "noscript"} and self.skip_depth:
            self.skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    @property
    def text(self) -> str:
        return clean_text("".join(self.parts))


def docx_texts(docx_path: Path) -> list[str]:
    doc = Document(docx_path)
    values: list[str] = []
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            values.append(text)
    for table in doc.tables:
        # python-docx exposes every grid slot of a merged cell as a Cell wrapper.
        # Count the underlying XML cell once so PDF-faithful rowspan/colspan output
        # is not falsely reported as underrepresenting duplicated export text.
        seen_cells: set[Any] = set()
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                for paragraph in cell.paragraphs:
                    text = clean_text(paragraph.text)
                    if text:
                        values.append(text)
    return values


def docx_image_count(docx_path: Path) -> int:
    doc = Document(docx_path)
    return len(doc.element.body.xpath(".//a:blip"))


def docx_table_count(docx_path: Path) -> int:
    return len(Document(docx_path).tables)


def html_visible_text(html_text: str) -> str:
    parser = VisibleTextParser()
    parser.feed(html_text)
    return parser.text


def strip_tags(value: str) -> str:
    return clean_text(re.sub(r"<[^>]+>", "", value))


def normalize_for_match(value: str) -> str:
    """Drop structural markers the generator legitimately rewrites, so canonical
    transformations (stripped 主标题： prefix, chapter numerals moved into the
    01-05 badge, { 标题 } braces, trailing colons) are not flagged as missing."""
    value = re.sub(r"^\s*(?:主标题|标题|文档标题|page\s*title)\s*[:：]\s*", "", value, flags=re.I)
    value = re.sub(r"[（(][一二三四五六七八九十0-9]+[）)]", "", value)
    value = re.sub(r"^[一二三四五六七八九十]+[、.．]", "", value)
    value = re.sub(r"^第[一二三四五六七八九十0-9]+[章节部分]", "", value)
    value = value.replace("{", "").replace("}", "")
    # A 「（…+X%）」title metric is split into the h2 + a separate green bar, and
    # other （…）asides may move; ignore parentheses when matching.
    value = re.sub(r"[（）()]", "", value)
    return re.sub(r"[：:\s]", "", value)


def count_class(html_text: str, class_name: str) -> int:
    pattern = rf'<[^>]+\bclass="[^"]*\b{re.escape(class_name)}\b[^"]*"'
    return len(re.findall(pattern, html_text, flags=re.I))


def structural_checks(html_text: str, visible: str, expected_images: int) -> tuple[dict[str, bool], list[str]]:
    h2_texts = [
        strip_tags(match)
        for match in re.findall(r"<h2\b[^>]*>(.*?)</h2>", html_text, flags=re.I | re.S)
    ]
    section_head_count = count_class(html_text, "section-head")
    spec_head_count = len(
        re.findall(
            r'<div\b[^>]*class="(?=[^"]*\bsection-head\b)(?=[^"]*\bspec-head\b)[^"]*"',
            html_text,
            flags=re.I,
        )
    )
    en_label_count = count_class(html_text, "en-label")
    label_line_count = count_class(html_text, "label-line")
    label_text_count = count_class(html_text, "label-text")
    label_plain_count = count_class(html_text, "label-plain")
    text_block_count = count_class(html_text, "text-block")
    caption_card_count = count_class(html_text, "caption-image-card")
    image_holder_count = count_class(html_text, "image-holder")
    # The hero-overlay background is an intentional replaceable layer injected by
    # the generator; it is NOT a DOCX content image and must never be wrapped in an
    # image-holder. Allow that single exception so a fully-correct page still passes.
    hero_overlay_count = len(re.findall(r"\bhero-overlay\b", html_text, flags=re.I))
    image_frame_count = count_class(html_text, "image-frame")
    screens_grid_count = count_class(html_text, "screens-grid")
    detail_screen_grid_count = count_class(html_text, "detail-screen-grid")
    sublevel_count = count_class(html_text, "sublevel")
    row_head_count = count_class(html_text, "row-head")
    metric_count = count_class(html_text, "metric-emphasis")
    video_case_count = count_class(html_text, "video-case-grid")
    video_case_card_count = count_class(html_text, "video-case-card")
    screen_count = len(re.findall(r"第\s*\d+\s*屏|【第\s*\d+\s*屏】", visible))
    # A genuine formula/caption under an image looks like "[图A] + [图B] = 效果":
    # bracketed terms joined by + AND resolved with an = result. Structured title
    # grammar such as "[品牌/系列] + 产品词 + 规格" also uses brackets and +, but is
    # NOT an image caption, so it must not trigger the caption-image-card rule.
    formula_caption_like = bool(
        re.search(r"\[[^\]]+\]\s*\+[^=\n]*=", visible)
    )
    overview_sublevels_needed = all(
        text in visible for text in ("【主图】", "首张主图：", "主图前5张：", "丰富素材图类型：")
    )
    sublevel_indent_present = all(
        token in html_text
        for token in (
            "--nested-group-indent: 42px",
            "--nested-text-offset: 25px",
            "margin-left: var(--nested-group-indent)",
            "padding-left: var(--nested-text-offset)",
        )
    )
    numbered_item_bolded = bool(
        re.search(r"<li\b[^>]*>\s*<b>\s*(?:\d+|[一二三四五六七八九十]+)[、.．]", html_text)
    )
    local_subtitle_plain = bool(
        re.search(
            r'class="[^"]*\blabel-plain\b[^"]*"[^>]*>\s*[（(][0-9一二三四五六七八九十]+[）)]',
            html_text,
        )
    )
    table_system_tokens = (
        "--table-font-size: 24px",
        "--table-header-weight: 700",
        "--table-body-weight: 400",
        "--table-body-bg: #f7f7f7",
        "--table-cell-radius: 10px",
        "--table-cell-gap: 8px",
        "--table-media-padding: 12px",
        "border-collapse: separate",
        "text-align: justify",
        "text-align-last: left",
    )

    checks = {
        "section_title_single_spaced_braces": bool(h2_texts)
        and all(re.fullmatch(r"\{ [^{}]+ \}", text) for text in h2_texts),
        "section_heads_include_spec_head_and_label": section_head_count > 0
        and section_head_count == spec_head_count == en_label_count,
        "label_lines_wrap_label_text": label_line_count == 0
        or label_line_count == label_text_count + label_plain_count,
        "white_text_blocks_used_for_label_groups": label_line_count < 6 or text_block_count > 0,
        "image_cards_use_image_holders": expected_images == 0
        or image_holder_count >= expected_images - hero_overlay_count,
        "caption_images_keep_text_below_images": not formula_caption_like or caption_card_count > 0,
        "many_screen_examples_use_detail_grid": screen_count < 5 or detail_screen_grid_count > 0,
        "detail_grid_not_four_column_screens_grid": screen_count < 5
        or screens_grid_count == 0,
        "overview_children_use_grey_sublevels": not overview_sublevels_needed
        or (sublevel_count >= 3 and sublevel_indent_present),
        "numbered_siblings_keep_equal_weight": not numbered_item_bolded,
        "local_numbered_subtitles_use_pink_marker": not local_subtitle_plain,
        "conversion_metrics_use_green_component": "效果数据" not in visible or metric_count > 0,
        "rounded_table_typography_system_present": all(token in html_text for token in table_system_tokens),
        "short_first_column_cells_use_row_headers": (
            "material-table" not in html_text and "spec-table" not in html_text
        ) or row_head_count > 0,
        "attribute_headers_default_to_red": count_class(html_text, "attr-head-gray") == 0
        or count_class(html_text, "is-semantic-alt") > 0,
        "video_case_headers_share_one_component": video_case_count == 0
        or (
            video_case_count == video_case_card_count
            and "video-demo" in html_text
            and "--video-header-height: 72px" in html_text
            and "height: var(--video-header-height)" in html_text
        ),
    }

    messages = {
        "section_title_single_spaced_braces": "Section h2 titles must be exactly `{ 标题 }`; double braces or missing inner spaces are not allowed.",
        "section_heads_include_spec_head_and_label": "Every section-head, including overview cards, must use spec-head and include the right INTRODUCTION label.",
        "label_lines_wrap_label_text": "Every label-line must contain label-text (colon title, pink bar) or label-plain (colon-less, no bar).",
        "white_text_blocks_used_for_label_groups": "Grey-panel label groups must be wrapped in white text-block modules instead of bare labels/lists.",
        "image_cards_use_image_holders": "Every DOCX image should be placed inside an image-holder so it keeps proportion and spacing.",
        "caption_images_keep_text_below_images": "Formula or caption text after an image must use caption-image-card, with image above and caption below.",
        "many_screen_examples_use_detail_grid": "Five or more screen/detail examples must use detail-screen-grid instead of a vertical list or cramped grid.",
        "detail_grid_not_four_column_screens_grid": "Screen/detail examples must not use the four-column screens-grid pattern; use the two-column detail-screen-grid.",
        "overview_children_use_grey_sublevels": "Children under a bracket parent such as 【主图】 must use grey-square sublevel items with the same 42px/25px indentation geometry as source-list children.",
        "numbered_siblings_keep_equal_weight": "Numbered siblings such as 1、/2、/3、 must keep equal weight; do not bold only the item that contains a colon.",
        "local_numbered_subtitles_use_pink_marker": "Module-local numbered subtitles must use label-text with the pink marker, not label-plain.",
        "conversion_metrics_use_green_component": "Effect/conversion metrics must use the green metric-emphasis component.",
        "rounded_table_typography_system_present": "Semantic tables must embed the canonical rounded 24px justified-body table-card system with equal image insets.",
        "short_first_column_cells_use_row_headers": "Short first-column body labels must use the bold row-head class.",
        "attribute_headers_default_to_red": "Attribute/keyword headers default to red; grey headers require an explicit is-semantic-alt role.",
        "video_case_headers_share_one_component": "视频案例 and 点击播放 must share the video-case-card component with equal-height headers and the play control.",
    }
    warnings = [message for key, message in messages.items() if not checks[key]]
    return checks, warnings


def validate(docx_path: Path, html_path: Path) -> dict[str, Any]:
    html_text = html_path.read_text(encoding="utf-8")
    visible = html_visible_text(html_text)
    source_texts = docx_texts(docx_path)
    expected_counts = Counter(source_texts)
    actual_counts = Counter()
    missing: list[str] = []
    underrepresented: list[dict[str, Any]] = []
    normalized_visible = normalize_for_match(visible)

    for text, expected in expected_counts.items():
        actual = visible.count(text)
        actual_counts[text] = actual
        if actual == 0:
            normalized_text = normalize_for_match(text)
            if normalized_text and normalized_text in normalized_visible:
                # Present after canonical rewrite (title prefix / chapter numeral
                # / brace / colon normalization); not a real omission.
                continue
            # A chapter title's 「（…X%）」metric is split into the h2 title plus a
            # green bar (separated by the INTRODUCTION label), so the full string is
            # never contiguous. Pass when the title core and the metric each appear.
            mt = re.match(r"^(.*?)[（(]\s*([^（）()]*\d+(?:\.\d+)?\s*%)\s*[)）]$", text)
            if mt:
                core_n = normalize_for_match(mt.group(1))
                metric_n = normalize_for_match(mt.group(2))
                if core_n and metric_n and core_n in normalized_visible and metric_n in normalized_visible:
                    continue
            # A video-play placeholder line OR a video URL under 主图视频 is
            # intentionally rendered as the 「点击播放」 play card, so its original
            # wording (the placeholder text or the raw link) is replaced by design.
            if "video-demo" in html_text and (
                re.search(r"视频播放按钮|播放按钮|主图视频示范|点击查看|点击播放|观看视频|视频入口", text)
                or re.search(r"https?://", text)
            ):
                continue
            # 一级模块标题在源文里写成 `1、主图规范`…`8、属性`，HTML 里按规范去掉了
            # `N、` 前缀渲染成 `{ 主图规范 }`。去掉阿拉伯/中文数字前缀后能在可见文本
            # 里找到，就是有意改写、不算缺字。
            stripped_prefix = re.sub(r"^\s*[（(]?[0-9一二三四五六七八九十]+[）)、.．]\s*", "", text)
            if stripped_prefix != text:
                stripped_n = normalize_for_match(stripped_prefix)
                if stripped_n and stripped_n in normalized_visible:
                    continue
            # A label may be split into separate modules (e.g. "2.优化前后图 商详转化率+2%"
            # becomes a "2.优化前后图" title plus a "商详转化率+2%" green metric bar).
            # Pass when every whitespace-separated chunk is individually present.
            chunks = [normalize_for_match(c) for c in text.split() if normalize_for_match(c)]
            if len(chunks) >= 2 and all(c in normalized_visible for c in chunks):
                continue
            missing.append(text)
        elif actual < expected:
            underrepresented.append({"text": text, "expected": expected, "actual": actual})

    expected_images = docx_image_count(docx_path)
    all_images = len(re.findall(r"<img\b", html_text, flags=re.I))
    # Synthetic Hero overlays are UI/runtime assets, not DOCX content images.
    synthetic_images = count_class(html_text, "hero-overlay")
    actual_images = max(0, all_images - synthetic_images)
    # A 模块化布局图 schematic is intentionally redrawn as a .module-layout block
    # instead of embedding the raw image, so each redraw stands in for one image.
    redrawn_images = len(re.findall(r'class="[^"]*module-layout', html_text))
    effective_images = actual_images + redrawn_images
    expected_tables = docx_table_count(docx_path)
    actual_table_like = (
        len(re.findall(r"<table\b", html_text, flags=re.I))
        + len(
            re.findall(
                r'class="[^"]*(?:word-table-spec|spec-table|ba-compare|video-case-grid|module-layout)',
                html_text,
            )
        )
    )

    css_checks = {
        "embedded_style": "<style>" in html_text,
        "embedded_font_face": "@font-face" in html_text,
        "hero_font": "JINGDONGLangZhengTi1-Bold" in html_text,
        "hero_mark_present": "OPERATION" in html_text and "STANDARDS" in html_text and "hero-mark" in html_text,
        "chinese_update_label": "更新日期" in html_text,
        "introduction_labels": "<strong>INTRODUCTION</strong>" in html_text,
        "panel_spacing_rule": ".gray-panel > * + *" in html_text,
        "single_file_images": "data:image/" in html_text if expected_images else True,
        "old_content_labels_absent": not any(
            value in html_text
            for value in [
                "MAIN IMAGE",
                "MAIN VIDEO",
                "LONG TITLE",
                "SHORT TITLE",
                "SELLING POINT",
                "PARAMETER",
                "DETAIL PAGE",
            ]
        ),
    }
    structure_checks, structure_warnings = structural_checks(html_text, visible, max(0, expected_images - redrawn_images))

    warnings: list[str] = []
    if missing:
        warnings.append("Some DOCX text fragments are missing from visible HTML text.")
    if underrepresented:
        warnings.append("Some repeated DOCX text fragments appear fewer times in HTML.")
    if effective_images != expected_images:
        warnings.append("HTML content image count plus semantic redraws must exactly match DOCX image occurrences.")
    if expected_tables and actual_table_like < expected_tables:
        warnings.append("DOCX contains tables; HTML may not preserve all table-like structures.")
    if not all(css_checks.values()):
        warnings.append("One or more required CSS or layout invariants are missing.")
    warnings.extend(structure_warnings)

    return {
        "source": str(docx_path),
        "html": str(html_path),
        "source_text_count": len(source_texts),
        "unique_source_text_count": len(expected_counts),
        "missing_text_count": len(missing),
        "missing_texts": missing[:50],
        "underrepresented_text_count": len(underrepresented),
        "underrepresented_texts": underrepresented[:50],
        "expected_image_count": expected_images,
        "actual_image_count": actual_images,
        "synthetic_image_count": synthetic_images,
        "redrawn_image_count": redrawn_images,
        "expected_table_count": expected_tables,
        "actual_table_like_count": actual_table_like,
        "css_checks": css_checks,
        "structure_checks": structure_checks,
        "passed": not warnings,
        "warnings": warnings,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate DOCX-to-HTML output.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("html", type=Path)
    parser.add_argument("--report", type=Path)
    parser.add_argument("--strict", action="store_true", help="Exit non-zero when validation warnings exist.")
    args = parser.parse_args()

    report = validate(args.docx, args.html)
    text = json.dumps(report, ensure_ascii=False, indent=2)
    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(text + "\n", encoding="utf-8")
    print(text)
    return 1 if args.strict and not report["passed"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
