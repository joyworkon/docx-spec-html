#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


def clean_text(value: str) -> str:
    value = re.sub(r"[\u200b-\u200f\ufeff]", "", value or "")
    return re.sub(r"\s+", " ", value).strip()


def paragraph_images(doc: DocumentObject, paragraph: Paragraph) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    for idx, blip in enumerate(paragraph._p.xpath(".//a:blip"), 1):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rid or rid not in doc.part.rels:
            continue
        rel = doc.part.rels[rid]
        target = rel.target_ref
        part = getattr(rel, "target_part", None)
        images.append(
            {
                "index_in_paragraph": idx,
                "relationship_id": rid,
                "target": target,
                "content_type": getattr(part, "content_type", None),
                "byte_size": len(part.blob) if part is not None and hasattr(part, "blob") else None,
            }
        )
    return images


def cell_manifest(doc: DocumentObject, cell: Any) -> dict[str, Any]:
    paragraphs: list[dict[str, Any]] = []
    for paragraph_index, paragraph in enumerate(cell.paragraphs):
        text = clean_text(paragraph.text)
        images = paragraph_images(doc, paragraph)
        if text or images:
            paragraphs.append(
                {
                    "paragraph_index": paragraph_index,
                    "text": text,
                    "style": paragraph.style.name if paragraph.style else "",
                    "images": images,
                }
            )
    return {
        "text": clean_text(" ".join(p.text for p in cell.paragraphs)),
        "paragraphs": paragraphs,
    }


def table_manifest(doc: DocumentObject, table: Table, block_index: int, table_index: int) -> dict[str, Any]:
    rows: list[dict[str, Any]] = []
    for row_index, row in enumerate(table.rows):
        rows.append(
            {
                "row_index": row_index,
                "cells": [cell_manifest(doc, cell) for cell in row.cells],
            }
        )
    return {
        "block_index": block_index,
        "type": "table",
        "table_index": table_index,
        "row_count": len(table.rows),
        "column_count": max((len(row.cells) for row in table.rows), default=0),
        "rows": rows,
    }


def extract_manifest(docx_path: Path) -> dict[str, Any]:
    doc = Document(docx_path)
    blocks: list[dict[str, Any]] = []
    paragraph_index = 0
    table_index = 0

    for block_index, child in enumerate(doc.element.body.iterchildren()):
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            text = clean_text(paragraph.text)
            images = paragraph_images(doc, paragraph)
            if text or images:
                blocks.append(
                    {
                        "block_index": block_index,
                        "type": "paragraph",
                        "paragraph_index": paragraph_index,
                        "text": text,
                        "style": paragraph.style.name if paragraph.style else "",
                        "images": images,
                        "image_count": len(images),
                        "is_heading_style": bool(paragraph.style and ("Heading" in paragraph.style.name or "标题" in paragraph.style.name)),
                    }
                )
            paragraph_index += 1
        elif isinstance(child, CT_Tbl):
            table_index += 1
            blocks.append(table_manifest(doc, Table(child, doc), block_index, table_index))

    all_texts: list[str] = []
    for block in blocks:
        if block["type"] == "paragraph" and block["text"]:
            all_texts.append(block["text"])
        elif block["type"] == "table":
            for row in block["rows"]:
                for cell in row["cells"]:
                    if cell["text"]:
                        all_texts.append(cell["text"])

    return {
        "source": str(docx_path),
        "block_count": len(blocks),
        "paragraph_count": paragraph_index,
        "table_count": table_index,
        "image_count": sum(len(block.get("images", [])) for block in blocks if block["type"] == "paragraph")
        + sum(
            len(paragraph["images"])
            for block in blocks
            if block["type"] == "table"
            for row in block["rows"]
            for cell in row["cells"]
            for paragraph in cell["paragraphs"]
        ),
        "visible_text_count": len([text for text in all_texts if text]),
        "title_candidate": next((block["text"] for block in blocks if block["type"] == "paragraph" and block["text"]), ""),
        "blocks": blocks,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract DOCX structure manifest for model-led HTML reconstruction.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()

    manifest = extract_manifest(args.docx)
    text = json.dumps(manifest, ensure_ascii=False, indent=2)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(text + "\n", encoding="utf-8")
    print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
