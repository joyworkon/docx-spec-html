#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import html
import json
import mimetypes
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from validate_output import validate


SKILL_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DESIGN = SKILL_ROOT / "references" / "design.md"
DEFAULT_FONT = SKILL_ROOT / "assets" / "fonts" / "JINGDONGLangZhengTi1-Bold.ttf"
DEFAULT_H2C = SKILL_ROOT / "assets" / "vendor" / "html2canvas.min.js"
DEFAULT_EDITOR = SKILL_ROOT / "assets" / "vendor" / "html-editor.html"


@dataclass
class ParagraphBlock:
    text: str
    images: list[str]
    style: str
    list_level: int | None = None  # Word numbering level (ilvl); None = not a list item


@dataclass
class TableBlock:
    table: Table


def clean_text(value: str) -> str:
    value = re.sub(r"[\u200b-\u200f\ufeff]", "", value or "")
    return re.sub(r"\s+", " ", value).strip()


def esc(value: str) -> str:
    return html.escape(value, quote=True)


def slugify(value: str) -> str:
    value = clean_text(value)
    value = re.sub(r"[\\/:*?\"<>|]+", "-", value)
    value = re.sub(r"\s+", "-", value).strip("-")
    return value[:80] or "docx"


def paragraph_list_level(paragraph: Paragraph) -> int | None:
    """Word numbering level (w:numPr/w:ilvl) for the paragraph, or None when the
    paragraph is not part of a list. This is the source's own hierarchy signal:
    None = heading/top level, 0 = first list level, 1+ = nested sub-levels."""
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        return 0
    try:
        return int(ilvl.get(qn("w:val")))
    except (TypeError, ValueError):
        return 0


def iter_blocks(doc: DocumentObject) -> list[ParagraphBlock | TableBlock]:
    blocks: list[ParagraphBlock | TableBlock] = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            text = clean_text(paragraph.text)
            images = paragraph_images(doc, paragraph)
            if text or images:
                blocks.append(ParagraphBlock(
                    text=text, images=images,
                    style=paragraph.style.name if paragraph.style else "",
                    list_level=paragraph_list_level(paragraph),
                ))
        elif isinstance(child, CT_Tbl):
            blocks.append(TableBlock(Table(child, doc)))
    return blocks


def paragraph_images(doc: DocumentObject, paragraph: Paragraph) -> list[str]:
    targets: list[str] = []
    for blip in paragraph._p.xpath(".//a:blip"):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid and rid in doc.part.rels:
            targets.append(doc.part.rels[rid].target_ref)
    return targets


def image_target_to_blob(doc: DocumentObject) -> dict[str, bytes]:
    blobs: dict[str, bytes] = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            blobs[rel.target_ref] = rel.target_part.blob
    return blobs


def extract_css(design_path: Path, font_path: Path | None) -> str:
    design = design_path.read_text(encoding="utf-8")
    if "## 十七、完整 CSS 模板" in design and "## 十八、验收与交付协议" in design:
        section = design.split("## 十七、完整 CSS 模板", 1)[1].split("## 十八、验收与交付协议", 1)[0]
    else:
        section = design
    blocks = re.findall(r"```css\n(.*?)```", section, flags=re.S)
    css = "\n\n".join(blocks)
    if not css.strip():
        raise ValueError(f"No CSS code block found in {design_path}")
    if font_path and font_path.exists():
        font_data = base64.b64encode(font_path.read_bytes()).decode("ascii")
        css = re.sub(
            r'src:\s*url\("JINGDONGLangZhengTi1-Bold\.ttf"\)\s*format\("truetype"\);',
            f'src: url("data:font/ttf;base64,{font_data}") format("truetype");',
            css,
        )
    return css + GENERIC_CSS


GENERIC_CSS = """

/* ===== Generic DOCX generator additions ===== */
body { background: #737373; }
.poster.auto-doc .hero { height: 500px; }
.poster.auto-doc .hero h1 { max-width: 760px; }
.poster.auto-doc .hero + .card { margin-top: -75px; }
.poster.auto-doc .gray-panel > * + * { margin-top: 18px; }
.poster.auto-doc .plain-block p {
  margin: 0;
  font-size: 19px;
  line-height: 1.55;
  font-weight: 400;
}
.poster.auto-doc .plain-block p + p { margin-top: 10px; }
.poster.auto-doc .image-holder {
  display: grid;
  place-items: center;
  min-height: 180px;
}
.poster.auto-doc .image-frame .image-holder { margin-top: 10px; }
.poster.auto-doc .wide-image .doc-image {
  width: auto;
  max-width: 100%;
  max-height: 560px;
  object-fit: contain;
}
.poster.auto-doc .doc-table-wrap {
  background: #fff;
  border-radius: 10px;
  padding: 18px;
}
.poster.auto-doc .doc-table {
  width: 100%;
  border-collapse: separate;
  border-spacing: 8px;
  table-layout: fixed;
}
.poster.auto-doc .doc-table th,
.poster.auto-doc .doc-table td {
  vertical-align: middle;
  border-radius: 8px;
  padding: 12px;
  font-size: 16px;
  line-height: 1.45;
}
.poster.auto-doc .doc-table th {
  background: #ff2b22;
  color: #fff;
  font-weight: 600;
  text-align: center;
}
.poster.auto-doc .doc-table td { background: #f7f7f7; }
.poster.auto-doc .doc-table .doc-image {
  width: auto;
  max-width: 100%;
  max-height: 300px;
  object-fit: contain;
}

/* Grey-square caption for example images (no red square, no highlight) */
.poster.auto-doc .caption-line {
  position: relative;
  margin: 0;
  padding-left: 25px;
  color: #555;
  font-size: 17px;
  line-height: 1.4;
  font-weight: 600;
}
.poster.auto-doc .caption-line::before {
  content: "";
  position: absolute;
  left: 0;
  top: 7px;
  width: 11px;
  height: 11px;
  border-radius: 3px;
  background: #d8d8d8;
}
.poster.auto-doc .text-block .caption-line { margin-top: 14px; }
.poster.auto-doc .text-block .image-holder { margin-top: 10px; }
/* Sub-level indent: example caption + images sit one level under their label,
   aligned with .source-list (28px). */
.poster.auto-doc .caption-line.indent { margin-left: 28px; }
.poster.auto-doc .image-holder.indent { margin-left: 28px; }
/* Third hierarchy level (Word ilvl>=1): deeper indent + hollow grey square. */
.poster.auto-doc .source-list li.deep { margin-left: 28px; }
.poster.auto-doc .source-list li.deep::before {
  background: transparent;
  border: 2px solid #c9c9c9;
}

/* Half-page-width images (used in 图文详情) */
.poster.auto-doc .half-image .doc-image {
  width: 100%;
  max-width: 600px;
  height: auto;
  object-fit: contain;
}

/* Example images under a label (e.g. 短标题「示例：」): half the container
   width and all equal width, regardless of each picture's aspect ratio. */
.poster.auto-doc .sample-image .doc-image {
  width: 50%;
  max-width: 50%;
  height: auto;
  object-fit: contain;
}

/* Module-layout schematic: a clean, watermark-free redraw of the
   "首张主图模块化布局图" reference image. The system is layout-AGNOSTIC: the
   container fixes only the yellow fill + red border; the actual block positions,
   sizes and wording come from the reference image. A model with vision must read
   the real image, place each .ml-block on the grid to mirror the source's row /
   column split and relative areas, and write the verbatim on-image text. */
/* 首张主图模块化布局重绘：黄底红框容器；内部用 .ml-grid 网格 + .ml-block 色块，
   每块的行列位置/跨度由模型按"那张图"的真实布局用内联 grid-column/grid-row 设定，
   不再写死京东模板。每块用一种明显区分的颜色（黄底/红框固定，其余各异），文字居中。 */
.poster.auto-doc .module-layout {
  max-width: 620px;        /* keep the source schematic's portrait proportion */
  margin: 0 auto;
  background: #fff8e1;     /* fixed light-yellow fill */
  border: 2px solid #ff2b22;  /* fixed red border */
  border-radius: 14px;
  padding: 16px;
}
/* The grid: the model sets grid-template-columns / -rows (and gap) inline to
   reproduce the reference image's split, e.g. style="grid-template-columns:1fr 2fr". */
.poster.auto-doc .module-layout .ml-grid {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 12px;
}
.poster.auto-doc .module-layout .ml-block {
  display: grid;
  place-items: center;
  text-align: center;       /* every block's text is centered */
  border-radius: 10px;
  padding: 18px 14px;
  line-height: 1.3;
  font-weight: 700;
  font-size: 22px;
}
.poster.auto-doc .module-layout .ml-lg { font-size: 34px; }   /* emphasize the main product */
.poster.auto-doc .module-layout .ml-tall { min-height: 280px; }
/* Distinct fills the model assigns so neighbouring blocks stay clearly separated. */
.poster.auto-doc .module-layout .ml-c1 { background: #cfe6ff; color: #134a73; }  /* blue */
.poster.auto-doc .module-layout .ml-c2 { background: #ffe0b3; color: #7a4a12; }  /* orange */
.poster.auto-doc .module-layout .ml-c3 { background: #d7f0d8; color: #1f5f2a; }  /* green */
.poster.auto-doc .module-layout .ml-c4 { background: #e7e0f5; color: #3a2f78; }  /* purple */
.poster.auto-doc .module-layout .ml-c5 { background: #ffd6e0; color: #8a2741; }  /* pink */
.poster.auto-doc .module-layout .ml-c6 { background: #fff3b0; color: #7a6512; }  /* amber */
.poster.auto-doc .module-layout .ml-c7 { background: #cdeeea; color: #0f5b54; }  /* teal */
.poster.auto-doc .module-layout .ml-c8 { background: #f6d9c0; color: #8a4a1f; }  /* tan */

/* 主图视频 play card — 「点击播放」 + play icon (solid dark-grey circle with a
   hollow knocked-out triangle), dark grey on a light-pink ground, centred. */
.poster.auto-doc .video-demo {
  display: flex;
  align-items: center;
  justify-content: center;
  gap: 18px;
  padding: 40px 24px;
  border-radius: 10px;
  background: #fbe2ec;
}
.poster.auto-doc .video-demo .vd-text {
  font-size: 30px;
  font-weight: 700;
  color: #555;
}
.poster.auto-doc .video-demo .vd-icon { width: 52px; height: 52px; display: block; }

/* Conversion-metric emphasis bar: white fill, green outline, centred.
   Black label + green "+X%" with an up-arrow. */
.poster.auto-doc .metric-emphasis {
  display: flex;
  align-items: center;
  justify-content: center;
  gap: 12px;
  width: 100%;
  margin: 0;
  padding: 18px 24px;
  border: 2px solid #47b250;
  border-radius: 14px;
  background: #fff;
  font-size: 24px;
  font-weight: 700;
}
.poster.auto-doc .metric-emphasis .metric-text {
  color: #111;
  line-height: 1;
}
.poster.auto-doc .metric-emphasis .metric-value {
  display: inline-flex;
  align-items: center;
  gap: 6px;
  color: #47b250;
  font-size: 40px;
  font-weight: 800;
  line-height: 1;
}
.poster.auto-doc .metric-emphasis .metric-arrow {
  height: 34px;
  width: auto;
  display: block;
}
/* Inside a white module, leave room between the bar and the table below it. */
.poster.auto-doc .text-block .metric-emphasis { margin: 0 0 12px; }
/* metric-emphasis multi-item: several indicators inline, separated by a left border */
.poster.auto-doc .metric-item { display: inline-flex; align-items: center; gap: 8px; }
.poster.auto-doc .metric-item + .metric-item { border-left: 1px solid #cfe8d4; padding-left: 12px; margin-left: 4px; }

/* 4-column optimisation matrix (优化内容/案例/优化前/优化后), first column groups rows via rowspan */
.poster.auto-doc .compare-matrix { width: 100%; border-collapse: collapse; table-layout: fixed; background: #fff; }
.poster.auto-doc .compare-matrix th, .poster.auto-doc .compare-matrix td { border: 1px solid #e5e5e5; padding: 12px; vertical-align: middle; text-align: center; }
.poster.auto-doc .compare-matrix thead th { background: #ff2b22; color: #fff; font-weight: 700; }
.poster.auto-doc .compare-matrix .cm-group { background: #f7f7f7; font-weight: 700; }
/* 图片单元格铺满：图片 td 去内边距，holder/图片撑满，与表格上下左右对齐。
   不用 object-fit/max-height——那会让图小一圈，且 html2canvas 不渲染 object-fit。 */
.poster.auto-doc .compare-matrix td.cm-img { padding: 0; }
.poster.auto-doc .compare-matrix td.cm-img .image-holder { margin: 0; width: 100%; min-height: 0; }
.poster.auto-doc .compare-matrix td.cm-img .doc-image { width: 100%; height: auto; display: block; border-radius: 10px; }
.poster.auto-doc .compare-matrix .doc-image { width: 100%; height: auto; display: block; }

/* material-type table (素材图类型/内容要求/示例), 「示例」header spans two image columns */
.poster.auto-doc .material-table { width: 100%; border-collapse: collapse; table-layout: fixed; background: #fff; }
.poster.auto-doc .material-table th, .poster.auto-doc .material-table td { border: 1px solid #e5e5e5;padding: 12px; vertical-align: middle; }
.poster.auto-doc .material-table thead th { background: #ff2b22; color: #fff; font-weight: 700; text-align: center; }
.poster.auto-doc .material-table .mt-type { text-align: center; font-weight: 700; }
.poster.auto-doc .material-table .mt-req { text-align: left; }
.poster.auto-doc .material-table .mt-eg { text-align: center; }
/* 示例图单元格铺满：图片 td 去内边距，holder/图片撑满，与表格上下左右对齐。
   不用 object-fit/max-height——那会让图小一圈，且 html2canvas 不渲染 object-fit。 */
.poster.auto-doc .material-table td.mt-eg { padding: 0; }
.poster.auto-doc .material-table td.mt-eg .image-holder { margin: 0; width: 100%; min-height: 0; }
.poster.auto-doc .material-table .mt-eg .doc-image { width: 100%; height: auto; display: block; border-radius: 10px; }
.poster.auto-doc .material-table col.mt-c1 { width: 16%; } .poster.auto-doc .material-table col.mt-c2 { width: 34%; }
.poster.auto-doc .material-table col.mt-c3 { width: 25%; } .poster.auto-doc .material-table col.mt-c4 { width: 25%; }

/* attribute / keyword-priority table: 2-col grey-red-grey text-only grid */
.poster.auto-doc .attr-table { width: 100%; border-collapse: collapse; table-layout: fixed; background: #fff; }
.poster.auto-doc .attr-table td { border: 1px solid #e5e5e5; padding: 14px; text-align: center; font-weight: 700; }
.poster.auto-doc .attr-table .at-grey { background: #f2f2f2; color: #333; }
.poster.auto-doc .attr-table .at-red { background: #ff2b22; color: #fff; }

/* tag-example image: doubled width (品质标签示例) */
.poster.auto-doc .tag-example-table .doc-image { width: 100%; max-height: 600px; object-fit: contain; }

/* layout placeholder for 首图模块化 schematic */
.poster.auto-doc .layout-holder { display: flex; align-items: center; justify-content: center; min-height: 200px; border: 2px dashed #d9d9d9; border-radius: 10px; background: #fafafa; overflow: hidden; }
.poster.auto-doc .layout-holder .doc-image { width: 100%; max-height: 520px; object-fit: contain; }


/* Before / after compare (优化前 grey, 优化后 red) */
.poster.auto-doc .ba-compare {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 14px;
}
.poster.auto-doc .ba-col {
  display: flex;
  flex-direction: column;
  gap: 10px;
}
.poster.auto-doc .ba-head {
  border-radius: 8px;
  padding: 10px;
  text-align: center;
  font-size: 18px;
  font-weight: 600;
}
.poster.auto-doc .ba-before { background: #f1f1f1; color: #555; }
.poster.auto-doc .ba-after { background: #ff2b22; color: #fff; }
.poster.auto-doc .ba-col .image-holder { min-height: 0; }
.poster.auto-doc .ba-text { margin: 0; font-size: 15px; line-height: 1.45; color: #333; }

/* Three-column spec table: 主图 narrowest, 内容要求 ~2x, 示例 half width */
.poster.auto-doc .spec-table {
  display: flex;
  flex-direction: column;
  gap: 8px;
}
.poster.auto-doc .spec-row {
  display: grid;
  grid-template-columns: 1fr 2fr 3fr;
  gap: 8px;
  align-items: stretch;
}
.poster.auto-doc .spec-cell {
  display: flex;
  flex-direction: column;
  justify-content: center;
  align-items: center;
  text-align: center;
  gap: 6px;
  background: #f7f7f7;
  border-radius: 8px;
  padding: 12px;
  font-size: 16px;
  line-height: 1.45;
  color: #333;
}
.poster.auto-doc .spec-head .spec-cell {
  background: #ff2b22;
  color: #fff;
  font-weight: 600;
  text-align: center;
  align-items: center;
}
.poster.auto-doc .spec-cell .image-holder {
  min-height: 0;
  width: 100%;
}
.poster.auto-doc .spec-cell .doc-image {
  /* Equal width via 100%, with height:auto driving the true aspect ratio.
     No max-height / object-fit: those force a wrong-ratio box that html2canvas
     (which ignores object-fit) would squish in the downloaded PNG. */
  display: block;
  width: 100%;
  height: auto;
}
.edit-toolbar {
  position: fixed;
  right: 18px;
  bottom: 18px;
  z-index: 9999;
  display: flex;
  gap: 8px;
  padding: 10px;
  border-radius: 10px;
  background: rgba(17, 17, 17, 0.88);
  color: #fff;
  font-family: "MiSans", "Microsoft YaHei", "PingFang SC", sans-serif;
}
.edit-toolbar button {
  border: 0;
  border-radius: 8px;
  padding: 8px 12px;
  background: #ff2b22;
  color: #fff;
  font-size: 13px;
  line-height: 1;
  cursor: pointer;
}
.edit-toolbar button.secondary { background: #555; }
.dl-page-btn {
  position: fixed;
  right: 18px;
  bottom: 18px;
  z-index: 9999;
  border: 0;
  border-radius: 10px;
  padding: 12px 18px;
  background: #ff2b22;
  color: #fff;
  font-family: "MiSans", "Microsoft YaHei", "PingFang SC", sans-serif;
  font-size: 14px;
  font-weight: 600;
  line-height: 1;
  cursor: pointer;
  box-shadow: 0 4px 14px rgba(0, 0, 0, 0.25);
}
.dl-page-btn[disabled] { opacity: 0.6; cursor: default; }
.edit-page-btn {
  position: fixed;
  right: 168px;      /* sit to the left of 下载整页图片 so they don't overlap */
  bottom: 18px;
  z-index: 9999;
  border: 0;
  border-radius: 10px;
  padding: 12px 18px;
  background: #1f2329;
  color: #fff;
  font-family: "MiSans", "Microsoft YaHei", "PingFang SC", sans-serif;
  font-size: 14px;
  font-weight: 600;
  line-height: 1;
  cursor: pointer;
  box-shadow: 0 4px 14px rgba(0, 0, 0, 0.25);
}
body.editing [contenteditable="true"] {
  outline: 2px dashed rgba(255, 43, 34, 0.75);
  outline-offset: 3px;
  cursor: text;
}

/* ===== Card body copy scaled ~1.5x (to match the target screenshot) =====
   Only the cards' body content grows; the hero and the title bar
   (.section-head: chapter number, { 标题 }, INTRODUCTION) keep their sizes.
   List dots, indents and module spacing scale with the text so proportions
   stay balanced. Higher specificity (.poster.auto-doc .class) overrides the
   §17 / base sizes above. */

/* 1) Body font sizes */
.poster.auto-doc .lead { font-size: 28px; }
.poster.auto-doc .plain-block p { font-size: 28px; }
.poster.auto-doc .red-list li,
.poster.auto-doc .red-list b,
.poster.auto-doc .red-list p { font-size: 28px; }
.poster.auto-doc .label-line { font-size: 28px; }
.poster.auto-doc .source-list li { font-size: 28px; }
.poster.auto-doc .source-list b { font-size: 28px; }
/* All text that follows a red-square title shares ONE size (28px): list items,
   sub-captions ("示例图：") and example lines. Table cells keep their own
   tighter tier below. */
.poster.auto-doc .example-line { font-size: 28px; }
.poster.auto-doc .caption-line { font-size: 28px; }
.poster.auto-doc .doc-table th,
.poster.auto-doc .doc-table td { font-size: 24px; }
.poster.auto-doc .spec-cell { font-size: 24px; }
.poster.auto-doc .ba-head { font-size: 27px; }
.poster.auto-doc .ba-text { font-size: 23px; }
/* Metric bar is NOT scaled with the body copy — it keeps the original
   24/40px text and 34px arrow defined above. */

/* Hero text scaled ~1.5x as well: main title, OPERATION STANDARDS, date.
   The hero grows taller so the enlarged title, rule and date keep their
   spacing and the white rule is not covered by the first card. */
.poster.auto-doc .hero { height: 600px; }
/* Clean hero background: flat brand-red, no gradient/grid texture/rings/dots. */
.poster.auto-doc .hero {
  background: #FF2B22;
}
.poster.auto-doc .hero::before,
.poster.auto-doc .robot-deco,
.poster.auto-doc .path-line,
.poster.auto-doc .rings,
.poster.auto-doc .ring { display: none; }
/* Replaceable hero background overlay: a transparent image layered ON TOP of the
   flat red but BEHIND the title text, and BLENDED with the red (mix-blend-mode)
   so a replaced image composites with the brand red instead of sitting opaquely
   on top. The red colour underneath is never changed.
   IMPORTANT: do NOT re-position the hero children here — the base CSS already
   gives .hero h1/.updated/.hero-mark/.hero-rule z-index:1, and .updated/.hero-mark
   are position:absolute (pinned to the corners). Overriding them to position
   relative drops the date/bracket into normal flow and breaks the hero. */
.poster.auto-doc .hero { position: relative; overflow: hidden; isolation: isolate; }
.poster.auto-doc .hero-overlay {
  position: absolute;
  inset: 0;
  width: 100%;
  height: 100%;
  object-fit: cover;
  object-position: center;
  z-index: 0;
  display: block;
  border: 0;
  mix-blend-mode: overlay;
}
/* Arrows rendered as inline SVG (not a CSS background image) so html2canvas
   keeps them in the downloaded full-page PNG. */
.poster.auto-doc .en-label span,
.poster.auto-doc .hero-rule { background: none; }
.poster.auto-doc .en-label span > svg,
.poster.auto-doc .hero-rule > svg { display: block; width: 100%; height: 100%; }
.poster.auto-doc .hero h1 { font-size: 102px; max-width: 1140px; }
.poster.auto-doc .hero-mark {
  font-size: 21px;
  width: 232px;
  height: 93px;
  border-radius: 48px;
  border-width: 2px;
}
.poster.auto-doc .updated { font-size: 27px; }
/* Bottom-align the update date with the white hero-rule arrow under the title. */
.poster.auto-doc .updated { bottom: 165px; }

/* 2) List dots scaled (~11 -> 16) */
.poster.auto-doc .red-list li::before,
.poster.auto-doc .label-line::before,
.poster.auto-doc .source-list li::before,
.poster.auto-doc .red-list li > p:not(.sub-dot)::before,
.poster.auto-doc .sub-dot::before {
  width: 16px;
  height: 16px;
  top: 13px;
  border-radius: 4px;
}
.poster.auto-doc .caption-line::before {
  width: 16px;
  height: 16px;
  top: 11px;
  border-radius: 4px;
}

/* 3) Indents scaled */
.poster.auto-doc .red-list li,
.poster.auto-doc .label-line,
.poster.auto-doc .caption-line { padding-left: 38px; }
.poster.auto-doc .source-list { margin-left: 42px; }
.poster.auto-doc .sub-dot { padding-left: 42px; }
.poster.auto-doc .caption-line.indent,
.poster.auto-doc .image-holder.indent { margin-left: 42px; }

/* 4) Module spacing scaled */
.poster.auto-doc .gray-panel > * + * { margin-top: 27px; }
.poster.auto-doc .text-block { padding: 33px 39px; }
.poster.auto-doc .lead { margin-bottom: 33px; padding: 36px 51px; }
.poster.auto-doc .red-list li { margin-bottom: 24px; }
.poster.auto-doc .spec-text .red-list li { margin-bottom: 27px; }
.poster.auto-doc .label-line { margin-bottom: 15px; }
.poster.auto-doc .source-list li { margin-top: 18px; }
/* Content split off after a label's colon: aligned under the title text, no red
   square, no pink bar, regular weight. */
.poster.auto-doc .label-rest {
  margin: 4px 0 0 38px;
  font-size: 28px;
  line-height: 1.5;
  font-weight: 400;
  color: #333;
}
/* A colon-less label keeps the red square but no pink highlight bar. */
.poster.auto-doc .label-plain { font-weight: 600; }
"""


EDITABLE_RUNTIME = """
<div class="edit-toolbar" data-html-edit-toolbar data-html2canvas-ignore>
  <button type="button" data-edit-toggle>编辑文字</button>
  <button type="button" data-edit-save>下载HTML</button>
  <button type="button" class="secondary" data-edit-exit>退出编辑</button>
</div>
<script>
(() => {
  const selectors = [
    "main.poster h1",
    "main.poster h2",
    "main.poster p",
    "main.poster li",
    "main.poster th",
    "main.poster td",
    "main.poster .label-text",
    "main.poster .chapter",
    "main.poster .en-label strong",
    "main.poster .hero-mark",
    "main.poster .image-caption-line span"
  ].join(",");
  const editableNodes = () => Array.from(document.querySelectorAll(selectors));
  const setEditing = (on) => {
    document.body.classList.toggle("editing", on);
    editableNodes().forEach((node) => {
      if (on) node.setAttribute("contenteditable", "true");
      else node.removeAttribute("contenteditable");
    });
    const toggle = document.querySelector("[data-edit-toggle]");
    if (toggle) toggle.textContent = on ? "正在编辑" : "编辑文字";
  };
  document.querySelector("[data-edit-toggle]")?.addEventListener("click", () => {
    setEditing(!document.body.classList.contains("editing"));
  });
  document.querySelector("[data-edit-exit]")?.addEventListener("click", () => setEditing(false));
  document.querySelector("[data-edit-save]")?.addEventListener("click", () => {
    const clone = document.documentElement.cloneNode(true);
    clone.querySelector("body")?.classList.remove("editing");
    clone.querySelectorAll("[contenteditable]").forEach((node) => node.removeAttribute("contenteditable"));
    const html = "<!doctype html>\\n" + clone.outerHTML;
    const blob = new Blob([html], { type: "text/html;charset=utf-8" });
    const link = document.createElement("a");
    link.href = URL.createObjectURL(blob);
    link.download = (document.title || "edited-output") + "-edited.html";
    link.click();
    URL.revokeObjectURL(link.href);
  });
})();
</script>
"""


def label_line(text: str) -> str:
    """A red-square label. A colon means the run up to (and incl.) the first colon
    is the TITLE — red square + pink highlight bar — and whatever follows the colon
    is split off below it with no highlight. With no colon anywhere, it is just a
    red square + plain text, no pink bar."""
    t = clean_text(text)
    cut = next((i + 1 for i, ch in enumerate(t) if ch in "：:"), -1)
    if cut != -1:
        prefix, rest = t[:cut], t[cut:].strip()
        html = f'<div class="label-line"><span class="label-text">{esc(prefix)}</span></div>'
        if rest:
            html += f'<div class="label-rest">{esc(rest)}</div>'
        return html
    return f'<div class="label-line"><span class="label-plain">{esc(t)}</span></div>'


def source_list(items: list) -> str:
    """items may be plain strings (grey level-0) or (level, text) tuples, where
    level>=1 renders as a deeper hollow-square sub-item."""
    rows = []
    for it in items:
        level, text = it if isinstance(it, tuple) else (0, it)
        if not clean_text(text):
            continue
        cls = ' class="deep"' if level and level >= 1 else ""
        rows.append(f"<li{cls}>{emphasize_prefix(esc(text))}</li>")
    return f'<ul class="source-list">{"".join(rows)}</ul>' if rows else ""


def emphasize_prefix(text: str) -> str:
    for mark in ("：", ":"):
        if mark in text and text.index(mark) <= 18:
            idx = text.index(mark) + 1
            return f"<b>{text[:idx]}</b>{text[idx:].strip()}"
    return text


def plain_block(items: list[str]) -> str:
    paras = "".join(f"<p>{esc(item)}</p>" for item in items if clean_text(item))
    return f'<div class="text-block plain-block">{paras}</div>' if paras else ""


def text_block(label: str, items: list[str]) -> str:
    return f'<div class="text-block">{label_line(label)}{source_list(items)}</div>'


def image_tag(target: str, blobs: dict[str, bytes], alt: str) -> str:
    data = blobs[target]
    mime = mimetypes.guess_type(target)[0] or "application/octet-stream"
    uri = f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"
    return f'<img class="doc-image" src="{uri}" alt="{esc(alt)}">'


def image_frame(label: str, target: str, blobs: dict[str, bytes]) -> str:
    safe_label = label or "示例图："
    return (
        '<div class="image-frame wide-image">'
        f"{label_line(safe_label)}"
        f'<div class="image-holder">{image_tag(target, blobs, safe_label)}</div>'
        "</div>"
    )


def render_table(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    rows_html: list[str] = []
    for row_idx, row in enumerate(table.rows):
        cells_html: list[str] = []
        tag = "th" if row_idx == 0 else "td"
        for cell in row.cells:
            texts = [clean_text(p.text) for p in cell.paragraphs if clean_text(p.text)]
            image_html = []
            for paragraph in cell.paragraphs:
                for target in paragraph_images(doc, paragraph):
                    if target in blobs:
                        image_html.append(f'<div class="image-holder">{image_tag(target, blobs, texts[0] if texts else "表格图片")}</div>')
            text_html = "<br>".join(esc(text) for text in texts)
            cells_html.append(f"<{tag}>{text_html}{''.join(image_html)}</{tag}>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return '<div class="doc-table-wrap"><table class="doc-table">' + "".join(rows_html) + "</table></div>"


def document_uses_heading_styles(blocks: list[ParagraphBlock | TableBlock]) -> bool:
    """True when the DOCX marks its chapters with real Word heading styles."""
    for block in blocks:
        if isinstance(block, ParagraphBlock) and block.text:
            style = block.style.lower()
            if style.startswith("heading") or "标题" in style:
                return True
    return False


CHAPTER_LABEL_SUFFIXES = ("展示", "示例", "布局", "流程", "说明", "组合")

# 商品信息运营规范的一级模块是固定且有限的这 8 个（顺序固定）。文档正文里
# 模块标题写成 `1、主图规�`…`8、属性`，正文要点也用 `1、2、3、` 编号——脚本
# 无法凭编号区分，故用白名单精确锚定：只有模块名命中白名单的 `N、xxx` 才是��节，
# 其余 `N、xxx`（如 `1、视频画面需清晰…`）一律留作正文，不升级为章节。
# 无编号短语（如 `卖点选词优先级`）也不得升级为第 9 个模块。
SPEC_MODULE_NAMES = (
    "主图规范", "主图视频", "长标题", "短标题",
    "通用卖点", "主推标签", "品质标签", "属性",
)
# 剥掉可能的 `N、`/`（一）`/`一、` 前缀后，与白名单精确比对用的正则
MODULE_NUM_PREFIX_RE = re.compile(r"^\s*[（(]?[0-9一二三四五六七八九十]+[）)、.．]\s*")


def spec_module_core(text: str) -> str:
    """剥掉一级编号前缀并去掉 （…X%）指标后，返回可与 SPEC_MODULE_NAMES 比对的核心词。"""
    core = MODULE_NUM_PREFIX_RE.sub("", text.strip())
    core = TITLE_METRIC_RE.sub("", core).strip()
    return core.rstrip("：:").strip()
TITLE_PAREN_RE = re.compile(r"[（(][^（）()]*[)）]\s*$")
TITLE_METRIC_RE = re.compile(r"[（(]\s*([^（）()]*?\d+(?:\.\d+)?\s*%)\s*[)）]\s*$")


def section_core(text: str) -> str:
    """Heading text with a trailing （…）parenthetical removed, for detection."""
    return TITLE_PAREN_RE.sub("", text).strip()


def section_title_metric(text: str) -> str | None:
    """A 「（…转化率提升+X%）」title suffix → its metric string for a green bar.
    A missing plus sign after 提升/增长 is normalised in so it renders green."""
    m = TITLE_METRIC_RE.search(text)
    if not m:
        return None
    return clean_text(m.group(1))  # keep source text verbatim (no injected +)


def is_section_heading(block: ParagraphBlock, heading_styles_present: bool = False) -> bool:
    text = block.text
    if not text or block.images:
        return False
    style = block.style.lower()
    if style.startswith("heading") or "标题" in style:
        return True
    if len(text) > 48:
        return False
    # A paragraph ending in a colon is a label/lead-in, never a chapter title.
    if text.endswith(("：", ":")):
        return False
    # 固定 8 模块白名单：剥掉 `N、`/`（一）` 前缀后精确命中白名单 → 一定是一级模块，
    # 即使带阿拉伯数字前缀（1、主图规范）也认，且优先于下面所有启发式。
    module_core = spec_module_core(text)
    if module_core in SPEC_MODULE_NAMES:
        return True
    # 「整体规范综述」这类概述卡保留识别。
    if section_core(text).startswith("整体规范"):
        return True
    # Explicit chapter numerals always win, including bracketed forms like （一）.
    if re.match(r"^[（(]?[一二三四五六七八九十]+[）)、.．]", text):
        return True
    if re.match(r"^第[一二三四五六七八九十0-9]+[章节部分]", text):
        return True
    if re.match(r"^【第[一二三四五六七八九十0-9-]+屏】", text):
        return True
    # 关键根治：阿拉伯数字 `N、xxx` 编号段落**不再**无条件升级为章节。
    # 商品信息运营规范里 `1、视频画面需清晰…`、`2、视频需搭配字幕…` 都是正文
    # 编号点，只有命中上面 8 模块白名单的 `N、模块名` 才是章节。此处直接不认，
    # 无编号短语（如 `卖点选词优先级`）也因不在白名单而留作模块内子项。
    return False


def clean_section_title(text: str) -> str:
    """Strip leading chapter numerals and trailing colons from a card title."""
    cleaned = text.strip()
    cleaned = re.sub(r"^[（(][一二三四五六七八九十0-9]+[）)]\s*", "", cleaned)
    cleaned = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", cleaned)
    # 也剥掉阿拉伯数字一级编号前缀（1、主图规范 → 主图规范），外层已有 01–08 章节号
    cleaned = re.sub(r"^[0-9]+[、.．]\s*", "", cleaned)
    cleaned = re.sub(r"^第[一二三四五六七八九十0-9]+[章节部分][:：、.．\s]*", "", cleaned)
    cleaned = TITLE_METRIC_RE.sub("", cleaned).strip()  # drop only a （…X%）metric, keep asides like （建议方向）
    cleaned = cleaned.rstrip("：:").strip()
    return cleaned or text.strip()


def strip_title_prefix(title: str) -> str:
    """Remove a leading label like 主标题：/ 标题: from the document main title."""
    return re.sub(r"^\s*(?:主标题|标题|文档标题|page\s*title)\s*[:：]\s*", "", title, flags=re.I)


SCREEN_LABEL_RE = re.compile(r"^第\s*[0-9一二三四五六七八九十]+(?:\s*[-–~至]\s*[0-9一二三四五六七八九十]+)?\s*屏")


def is_label(text: str) -> bool:
    if not text or len(text) > 40:
        return False
    # A standalone "XXX：" label, allowing a trailing aside like 卖点示例组合：（无优先级）.
    if re.sub(r"[（(][^（）()]*[)）]\s*$", "", text).strip().endswith(("：", ":")):
        return True
    if re.match(r"^[0-9]{1,2}[.、．]", text):
        # A numbered line is a label only when it's a title (no sentence
        # punctuation). A numbered sentence ("1、商品品牌…一致。") is a list item.
        body = re.sub(r"^[0-9]{1,2}[.、．]\s*", "", text)
        return not re.search(r"[，。；！？,;]", body)
    if SCREEN_LABEL_RE.match(text):
        return True
    return False


def is_bare_heading(text: str) -> bool:
    """A short standalone heading with no colon and no sentence punctuation, e.g.
    補充規則 — it opens a red-square container for the deeper lines beneath it."""
    t = clean_text(text)
    return bool(t) and len(t) <= 10 and not re.search(r"[：:，。；！？,.!?]", t)


def looks_like_image_title(text: str) -> bool:
    """A short noun-phrase line that titles the picture(s) right below it, e.g.
    搜索列表页展示 / 参数楼层商详页展示 / 商品参数展示 / 示例图。Used to lift such a line
    into the image's red-square label instead of leaving it as loose body text."""
    t = clean_text(text)
    if not t or len(t) > 16 or re.search(r"[，。！？,!?]", t):
        return False
    return t.endswith(("展示", "示例", "示范", "图", "流程", "说明")) or t.endswith(("：", ":")) or len(t) <= 8


def split_sections(blocks: list[ParagraphBlock | TableBlock]) -> tuple[str, list[tuple[str, list[ParagraphBlock | TableBlock]]]]:
    first_text_index = next((i for i, block in enumerate(blocks) if isinstance(block, ParagraphBlock) and block.text), None)
    if first_text_index is None:
        return "未命名规范", [("整体规范综述", blocks, None)]
    title = strip_title_prefix(blocks[first_text_index].text)  # type: ignore[union-attr]
    body = blocks[first_text_index + 1 :]
    heading_styles_present = document_uses_heading_styles(blocks)
    sections: list[tuple[str, list[ParagraphBlock | TableBlock], str | None]] = []
    current_title = "整体规范综述"
    current_metric: str | None = None
    current_blocks: list[ParagraphBlock | TableBlock] = []

    for block in body:
        if isinstance(block, ParagraphBlock) and is_section_heading(block, heading_styles_present):
            if current_blocks:
                sections.append((current_title, current_blocks, current_metric))
            current_title = clean_section_title(block.text)
            current_metric = section_title_metric(block.text)
            current_blocks = []
        elif (
            isinstance(block, ParagraphBlock)
            and block.text
            and not block.images
            and block.text.rstrip("：:").strip() == current_title.rstrip("：:").strip()
        ):
            # Drop a paragraph that merely restates the current card title
            # (e.g. an "整体规范综述：" lead-in under the 整体规范综述 card).
            continue
        else:
            current_blocks.append(block)
    if current_blocks:
        sections.append((current_title, current_blocks, current_metric))
    if not sections:
        sections.append(("整体规范综述", body, None))
    return title, sections


BRACKET_RE = re.compile(r"^\s*(【[^】]+】)\s*(.*)$")
CIRCLED_RE = re.compile(r"^\s*[①②③④⑤⑥⑦⑧⑨⑩⑪⑫]")
CONVERSION_RE = re.compile(r"^[一-龥A-Za-z·]{2,12}\s*[+＋]\s*\d+(?:\.\d+)?\s*%$")
# "前缀：内容" lead-in, e.g. 总结：…/字数范围：…/卖点建议顺序：… — gets a red square + pink highlight.
COLON_PREFIX_RE = re.compile(r"^([^：:\n]{1,18}[：:])(.+)$")
# A conversion metric embedded inside a longer label, e.g. "2.优化前后图 商详转化率+2%".
METRIC_INLINE_RE = re.compile(r"[一-龥A-Za-z·]{2,12}\s*[+＋]\s*\d+(?:\.\d+)?\s*%")


def split_label_metric(label: str) -> tuple[str | None, str | None]:
    """Pull an embedded "XX率+X%" metric out of a label so it can be rendered as a
    standalone green emphasis bar, leaving the rest of the label as the title."""
    match = METRIC_INLINE_RE.search(label)
    if not match:
        return label, None
    metric = clean_text(match.group(0))
    rest = clean_text(label[: match.start()] + " " + label[match.end():]).strip(" 　：:")
    return (rest or None), metric


def is_conversion_metric(text: str) -> bool:
    """A standalone "XX率+X%" style metric line that deserves green emphasis."""
    return bool(CONVERSION_RE.fullmatch(clean_text(text)))


METRIC_ARROW_SVG = (
    '<svg class="metric-arrow" xmlns="http://www.w3.org/2000/svg" '
    'viewBox="0 0 32.34917 40.82425" fill="none" aria-hidden="true">'
    '<path d="M16.405537,0L1.1559057,14.706532L11.676984,15.342317Q11.316808,35.421619,0,40.824245'
    'Q19.693825,39.944675,21.828087,15.110984L32.349167,15.74677L16.405537,0Z" fill="#47B250"/></svg>'
)

# The connected line+hook arrows rendered as INLINE svg (not CSS background) so
# html2canvas keeps them in the downloaded PNG. en-label = red, pointing left
# (flipped); hero-rule = white, pointing right.
_ARROW_PATH = (
    "M130.44582 9.5L0 9.5L0 12.5L138 12.5L138.20905 9.5146379Q137.62546 9.4325085 136.69609 9.1682091"
    "Q134.81464 8.6331568 133.25896 7.7437816Q128.52896 5.0396996 128.52898 -2.420493e-06"
    "L125.52898 2.420493e-06Q125.52896 3.7219667 127.51295 6.5661678Q128.67047 8.2256107 130.44582 9.5Z"
)
EN_LABEL_ARROW_SVG = (
    '<svg viewBox="0 0 138.209 12.5" preserveAspectRatio="none" xmlns="http://www.w3.org/2000/svg" '
    'aria-hidden="true"><g transform="translate(138.209,0) scale(-1,1)">'
    f'<path d="{_ARROW_PATH}" fill="#ff2b22"/></g></svg>'
)
HERO_RULE_ARROW_SVG = (
    '<svg viewBox="0 0 138.209 12.5" preserveAspectRatio="none" xmlns="http://www.w3.org/2000/svg" '
    f'aria-hidden="true"><path d="{_ARROW_PATH}" fill="#ffffff"/></svg>'
)

# Fully transparent 1x1 PNG: the default hero-overlay so the delivered page shows
# clean red. Replace this image in the editor (双击 → 替换图片) to float any picture
# ON TOP of the red without touching the red background colour underneath.
HERO_OVERLAY_SRC = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
)


def metric_emphasis(text: str) -> str:
    cleaned = clean_text(text)
    # Label stays black; the "+X%" (sign included) goes green with an up-arrow.
    match = re.match(r"^(.*?)([+＋]?\s*\d+(?:\.\d+)?\s*%)$", cleaned)
    if match and match.group(1).strip():
        head, value = match.group(1).strip(), match.group(2)
        inner = (
            f'<span class="metric-text">{esc(head)}</span>'
            f'<span class="metric-value">{esc(value)}{METRIC_ARROW_SVG}</span>'
        )
    else:
        inner = f'<span class="metric-text">{esc(cleaned)}</span>'
    return f'<div class="metric-emphasis">{inner}</div>'


# Strong "this is a clickable video" signals — trigger the 点击播放 card anywhere.
VIDEO_DEMO_RE = re.compile(r"主图视频示范|视频示范|示范视频|视频演示|视频播放按钮|播放按钮")
# Weak signals — a link or a "click to view" hint. These are only treated as a
# video play card INSIDE a 主图视频 section, so they don't hijack other sections.
VIDEO_HINT_RE = re.compile(r"查看链接|视频链接|视频地址|点击查看|点击播放|观看视频|视频入口|扫码观看|视频如下")
# A video link (http/https) provided under the 主图视频 section: its title + link
# are replaced by the 点击播放 card.
VIDEO_URL_RE = re.compile(r"https?://\S+")


def video_demo_box() -> str:
    """A 主图视频 play card: bold dark-grey 「点击播放」 label + a play icon (a solid
    dark-grey circle with a knocked-out / hollow triangle), centred on a light-pink
    ground. Triggered by a video-play placeholder line OR a video URL under the
    主图视频 card."""
    icon = (
        '<svg class="vd-icon" viewBox="0 0 48 48" xmlns="http://www.w3.org/2000/svg" aria-hidden="true">'
        # evenodd: full disc minus a triangle subpath -> the triangle is a hole.
        '<path fill-rule="evenodd" fill="#555" d="M24 1A23 23 0 1 0 24 47A23 23 0 1 0 24 1Z'
        'M19 14.5L35 24L19 33.5Z"/></svg>'
    )
    return f'<div class="video-demo"><span class="vd-text">点击播放</span>{icon}</div>'


def lead_block(text: str) -> str:
    return f'<p class="lead">{esc(text)}</p>'


def caption_line(text: str) -> str:
    """Grey-square caption (no red square, no highlight) for example images."""
    return f'<div class="caption-line">{esc(text)}</div>'


def red_list_block(items: list[str]) -> str:
    rows = []
    for item in items:
        item = clean_text(item)
        if not item:
            continue
        bracket = BRACKET_RE.match(item)
        if bracket:
            rows.append(f"<li><b>{esc(bracket.group(1))}</b>{esc(bracket.group(2))}</li>")
            continue
        colon = COLON_PREFIX_RE.match(item)
        if colon:
            rows.append(f"<li><b>{esc(colon.group(1))}</b>{esc(colon.group(2).strip())}</li>")
            continue
        rows.append(f"<li>{esc(item)}</li>")
    if not rows:
        return ""
    return f'<div class="text-block"><ul class="red-list">{"".join(rows)}</ul></div>'


def module_layout(items: list[str], fallback: str) -> str:
    """No-vision FALLBACK redraw of a「首张主图模块化布局图」schematic.

    This runs only when no model vision is available. It deliberately does NOT
    invent a layout (no "店铺名称" top row, no default "商品主图" / "质保承诺"
    slots — those guesses misled real models into copying a template instead of
    the actual image). It just stacks the REAL captured 主图首张 module names in a
    single column, verbatim, each in a distinct colour. An agent WITH image-reading
    ability must replace this with a faithful redraw that mirrors the real image's
    rows/columns and proportions (see references/design.md).

    Falls back to the raw image when fewer than two real module names are known."""
    names = [re.split(r"[：:]", it, 1)[0].strip() for it in items]
    names = [n for n in names if n]
    if len(names) < 2:
        return fallback

    palette = ["ml-c1", "ml-c2", "ml-c3", "ml-c4", "ml-c5", "ml-c6", "ml-c7", "ml-c8"]
    blocks = "".join(
        f'<div class="ml-block {palette[i % len(palette)]}">{esc(n)}</div>'
        for i, n in enumerate(names)
    )
    # Single-column stack of the real names — honest about what is known, with no
    # fabricated structure. The vision-capable redraw supplies the true grid.
    return (
        '<div class="module-layout"><div class="ml-grid" '
        'style="grid-template-columns:1fr;">' + blocks + "</div></div>"
    )


GENERIC_EXAMPLE_RE = re.compile(r"^示例图?\s*$")


def is_generic_example_label(label: str | None) -> bool:
    """A bare 示例 / 示例图 with NO colon and no title of its own → grey caption.
    A colon form (示例图：/ 示例：) is a real label-line and stays red."""
    return bool(label) and bool(GENERIC_EXAMPLE_RE.match(label.strip()))


def grouped_text_block(label: str | None, items: list[str], images: list[str], blobs: dict[str, bytes], half: bool, metric: str | None = None) -> str:
    """One white module that merges a label with its sub-items and example images.
    An embedded "XX率+X%" metric (pulled out of the label) renders as a green bar
    directly under the title, so it spans the module's inner width.

    Image caption rules: never synthesise a "示例图：" title. Use the source's own
    label as the picture's title — a real title (搜索列表页展示, 示例图：…) renders as a
    red-square label-line; a bare 示例/示例图 degrades to a grey caption-line."""
    generic = is_generic_example_label(label)
    inner = ""
    if label and generic:
        inner += caption_line(clean_text(label))
    elif label:
        inner += label_line(label)
    if metric:
        inner += metric_emphasis(metric)
    inner += source_list(items)
    real_images = [t for t in images if t in blobs]
    if real_images:
        # Indent the pictures one level (28px) only under a real red-square label,
        # so they align with .source-list sub-items; grey captions don't indent.
        indent = " indent" if (label and not generic) else ""
        if half:
            holder_class = "image-holder half-image" + indent
        elif len(real_images) >= 2:
            # Multiple stacked reference images: half container width, all equal.
            holder_class = "image-holder sample-image" + indent
        else:
            holder_class = "image-holder" + indent
        alt = clean_text(label) if label else "示意图"
        for target in real_images:
            inner += f'<div class="{holder_class}">{image_tag(target, blobs, alt)}</div>'
    return f'<div class="text-block">{inner}</div>' if inner else ""


def classify_table(table: Table) -> str:
    if not table.rows:
        return "generic"
    header = " ".join(clean_text(cell.text) for cell in table.rows[0].cells)
    ncol = len(table.rows[0].cells)
    if "优化前" in header or "优化后" in header:
        return "before_after"
    if ncol >= 3 and ("内容要求" in header or "示例" in header):
        return "spec"
    if ncol == 2:
        return "before_after"
    return "generic"


def before_after(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    rows = list(table.rows)
    if not rows:
        return ""
    header = [clean_text(cell.text) for cell in rows[0].cells]
    cols: list[str] = []
    for ci in range(len(rows[0].cells)):
        head = header[ci] if ci < len(header) else ""
        is_before = "前" in head or (ci == 0 and "后" not in head)
        head_class = "ba-before" if is_before else "ba-after"
        body = ""
        for row in rows[1:]:
            if ci >= len(row.cells):
                continue
            cell = row.cells[ci]
            for paragraph in cell.paragraphs:
                for target in paragraph_images(doc, paragraph):
                    if target in blobs:
                        body += f'<div class="image-holder">{image_tag(target, blobs, head)}</div>'
                txt = clean_text(paragraph.text)
                if txt:
                    body += f'<p class="ba-text">{esc(txt)}</p>'
        cols.append(f'<div class="ba-col"><div class="ba-head {head_class}">{esc(head)}</div>{body}</div>')
    return f'<div class="ba-compare">{"".join(cols)}</div>'


def spec_table(table: Table, doc: DocumentObject, blobs: dict[str, bytes]) -> str:
    rows_html: list[str] = []
    for row_idx, row in enumerate(table.rows):
        cells_html: list[str] = []
        for cell in row.cells:
            content = ""
            for paragraph in cell.paragraphs:
                txt = clean_text(paragraph.text)
                if txt:
                    content += f"<span>{esc(txt)}</span>"
                for target in paragraph_images(doc, paragraph):
                    if target in blobs:
                        content += f'<div class="image-holder">{image_tag(target, blobs, "示例")}</div>'
            cells_html.append(f'<div class="spec-cell">{content}</div>')
        row_class = "spec-row spec-head" if row_idx == 0 else "spec-row"
        rows_html.append(f'<div class="{row_class}">{"".join(cells_html)}</div>')
    return f'<div class="spec-table">{"".join(rows_html)}</div>'


def table_group(label: str | None, table: Table, doc: DocumentObject, blobs: dict[str, bytes], metric: str | None = None) -> str:
    kind = classify_table(table)
    if kind == "before_after":
        inner = before_after(table, doc, blobs)
    elif kind == "spec":
        inner = spec_table(table, doc, blobs)
    else:
        label_html = label_line(label) if label else ""
        return f'<div class="text-block">{label_html}</div>{render_table(table, doc, blobs)}' if label else render_table(table, doc, blobs)
    label_html = label_line(label) if label else ""
    # Embedded metric (e.g. 商详转化率+2%) sits under the title, spanning the same
    # width as the before/after (优化前+优化后) columns below it.
    metric_html = metric_emphasis(metric) if metric else ""
    return f'<div class="text-block">{label_html}{metric_html}{inner}</div>'


def render_section_blocks(blocks: list[ParagraphBlock | TableBlock], doc: DocumentObject, blobs: dict[str, bytes], *, is_intro: bool = False, half_images: bool = False, video_section: bool = False) -> str:
    rendered: list[str] = []
    plain_items: list[str] = []
    bracket_items: list[str] = []
    pending_label: str | None = None
    pending_metric: str | None = None
    pending_items: list[str] = []
    pending_images: list[str] = []
    module_items: list[str] = []  # captured 主图首张 module names, for the layout redraw
    lead_done = False
    video_card_done = False  # in a 主图视频 section, emit the play card only once

    def flush_plain() -> None:
        nonlocal plain_items
        if plain_items:
            rendered.append(plain_block(plain_items))
            plain_items = []

    def flush_bracket() -> None:
        nonlocal bracket_items
        if bracket_items:
            rendered.append(red_list_block(bracket_items))
            bracket_items = []

    def flush_label() -> None:
        nonlocal pending_label, pending_metric, pending_items, pending_images, module_items
        if pending_label is None and not pending_images and not pending_metric and not pending_items:
            return
        item_texts = [it[1] if isinstance(it, tuple) else it for it in pending_items]
        # Capture the 主图首张 module list so a later 模块化布局图 can be redrawn.
        if pending_label and "主图首张" in pending_label:
            mods = [t for t in item_texts if ("：" in t or ":" in t)]
            if mods:
                module_items = mods
        # Redraw "首张主图模块化布局图" as a clean module schematic (no watermark)
        # instead of embedding the raw reference screenshot.
        if pending_label and "模块化布局图" in pending_label and pending_images and module_items:
            real = [t for t in pending_images if t in blobs]
            fallback = "".join(
                f'<div class="image-holder">{image_tag(t, blobs, "模块化布局图")}</div>' for t in real
            )
            ml = module_layout(module_items, fallback)
            rendered.append(f'<div class="text-block">{label_line(pending_label)}{ml}</div>')
        else:
            rendered.append(grouped_text_block(pending_label, pending_items, pending_images, blobs, half_images, pending_metric))
        pending_label = None
        pending_metric = None
        pending_items = []
        pending_images = []

    def deeper_follows(idx: int, cur_level: int | None) -> bool:
        """True when the next content line is one hierarchy step deeper than the
        current line — i.e. the current line is a heading that opens a grey
        container. A deeper Word list level, a manual ①②③, or an image all count;
        a same/shallower level or an end means the current line is a sibling item."""
        base = -1 if cur_level is None else cur_level
        for nb in blocks[idx + 1:]:
            if isinstance(nb, TableBlock):
                return False
            if not nb.text:
                if nb.images:
                    return True
                continue
            if CIRCLED_RE.match(nb.text):
                return True
            return nb.list_level is not None and nb.list_level > base
        return False

    def table_follows(idx: int) -> bool:
        """True when the next non-empty block is a Word table — used to lift a
        caption-like line (e.g. 完整主图结构规范示例) into that table's title."""
        for nb in blocks[idx + 1:]:
            if isinstance(nb, TableBlock):
                return True
            if isinstance(nb, ParagraphBlock):
                if nb.text or nb.images:
                    return False
                continue
            return False
        return False

    for idx, block in enumerate(blocks):
        if isinstance(block, TableBlock):
            flush_plain()
            flush_bracket()
            if pending_label is not None and not pending_items and not pending_images:
                label = pending_label
                metric = pending_metric
                pending_label = None
                pending_metric = None
                rendered.append(table_group(label, block.table, doc, blobs, metric=metric))
            else:
                flush_label()
                rendered.append(table_group(None, block.table, doc, blobs))
            continue

        if block.images and not block.text:
            flush_bracket()
            # If the picture is directly preceded by a short title line (展示/示例/
            # 图…), lift it out of the loose text into the image's red-square label.
            if pending_label is None and not pending_images and plain_items and looks_like_image_title(plain_items[-1]):
                pending_label = plain_items.pop()
            flush_plain()
            # Accumulate into pending_images so consecutive image-only paragraphs
            # land in ONE module and share a single uniform width class.
            pending_images.extend(block.images)
            continue

        text = block.text
        if not text:
            continue

        if is_conversion_metric(text):
            flush_plain()
            flush_bracket()
            flush_label()
            rendered.append(metric_emphasis(text))
            continue

        is_video_strong = bool(VIDEO_DEMO_RE.search(text))
        is_video_weak = video_section and not is_video_strong and bool(
            VIDEO_URL_RE.search(text) or VIDEO_HINT_RE.search(text)
        )
        if is_video_strong or is_video_weak:
            flush_plain()
            flush_bracket()
            if is_video_weak:
                # The card replaces the title + link/hint: drop a pending bare title
                # (e.g. 主图视频：/ 视频链接：/ 查看链接) instead of rendering it.
                pending_label = None
                pending_metric = None
                pending_items = []
                pending_images = []
            else:
                flush_label()
            # In a 主图视频 section collapse repeated placeholders/links to one card.
            if not (video_section and video_card_done):
                rendered.append(video_demo_box())
                video_card_done = True
            continue

        # A caption-like line immediately followed by a TABLE is that table's
        # title: lift it to a red-square label-line in the SAME module as the
        # table (e.g. 完整主图结构规范示例 above the 主图/内容要求/示例 spec table),
        # even when a previous container is still open. Without this the line is
        # swallowed as a grey sub-item of the open container and the table loses
        # its title.
        if looks_like_image_title(text) and table_follows(idx):
            flush_plain()
            flush_bracket()
            flush_label()
            pending_label = text
            pending_items = []
            pending_images = []
            continue

        if BRACKET_RE.match(text):
            flush_plain()
            flush_label()
            bracket_items.append(text)
            continue

        # A colon label opens a new red container — UNLESS it is itself a Word list
        # item nested inside an already-open container (e.g. 性能及其他利益点： at
        # ilvl=0 under 2.详细规范：), in which case it stays a grey sub-item below.
        if is_label(text) and not (pending_label is not None and block.list_level is not None):
            flush_plain()
            flush_bracket()
            flush_label()
            rest, metric = split_label_metric(text)
            pending_label = rest
            pending_metric = metric
            pending_items = []
            pending_images = []
            continue

        # A short noun-phrase line that titles the picture(s) right below it
        # (搜索列表页展示 / 参数楼层商详页展示 / 商品参数展示) opens a fresh red-square
        # image module, breaking any container or image group currently open.
        if block.list_level is None and looks_like_image_title(text) and deeper_follows(idx, None):
            flush_plain()
            flush_bracket()
            flush_label()
            pending_label = text
            pending_items = []
            pending_images = []
            continue

        # Inside an open red container, every line is a grey sub-detail whose
        # indent comes straight from its Word level (ilvl); an ilvl=None line sits
        # at level 0.
        if pending_label is not None:
            level = block.list_level if block.list_level is not None else 0
            pending_items.append((level, text))
            continue

        # No container open: this line is top-level.
        flush_label()
        if is_intro and not lead_done and not rendered and not plain_items and not bracket_items:
            rendered.append(lead_block(text))
            lead_done = True
        elif (block.list_level is None and (is_bare_heading(text) or COLON_PREFIX_RE.match(text))
              and deeper_follows(idx, None)):
            # An ilvl=None heading (補充規則, or 前缀：… with deeper lines under it)
            # opens a red container; the deeper lines become grey sub-items.
            flush_plain()
            flush_bracket()
            pending_label = text
            pending_items = []
            pending_images = []
        elif COLON_PREFIX_RE.match(text) or block.list_level is not None:
            # A standalone "前缀：内容" line, OR any Word-bulleted line sitting
            # directly under the chapter (e.g. 参数楼层&属性 的前两段) → red-square
            # item. No colon ⇒ red square only, no pink bar.
            flush_plain()
            bracket_items.append(text)
        else:
            # A plain unbulleted paragraph with no heading role → its own white box.
            flush_bracket()
            plain_items.append(text)

    flush_plain()
    flush_bracket()
    flush_label()
    return "".join(rendered) or plain_block([""])


def hero_title_html(title: str) -> str:
    """Force the "商品信息运营规范" category suffix onto the hero title's second line."""
    title = strip_title_prefix(title)
    marker = "商品信息运营规范"
    if marker in title:
        prefix = title[: title.index(marker)].rstrip()
        return f"{esc(prefix)}<br>{esc(marker)}" if prefix else esc(marker)
    return esc(title)


def section_head(num: int | None, title: str) -> str:
    chapter = f'<span class="chapter">{num:02d}</span>' if num is not None else ""
    return (
        '<div class="section-head spec-head">'
        f"{chapter}"
        f"<h2>{{ {esc(title)} }}</h2>"
        f'<div class="en-label"><strong>INTRODUCTION</strong><span>{EN_LABEL_ARROW_SVG}</span></div>'
        "</div>"
    )


def normalize_update_label(value: str | None, docx_path: Path | None = None) -> str:
    if not value:
        # Default to the source file's last-modified date, to the day.
        d = date.fromtimestamp(docx_path.stat().st_mtime) if docx_path is not None else date.today()
        return f"更新日期 {d.year}年{d.month}月{d.day}日"
    text = clean_text(value)
    if re.search(r"[\u4e00-\u9fff]", text):
        return text
    match = re.search(r"(20\d{2})[.\-/年 ]+(\d{1,2})(?:[.\-/月 ]+(\d{1,2}))?", text)
    if match:
        year, month, day = match.group(1), int(match.group(2)), match.group(3)
        if day:
            return f"更新日期 {year}年{month}月{int(day)}日"
        return f"更新日期 {year}年{month}月"
    return f"更新日期 {text}"


def render_card(title: str, body: str, num: int | None, metric: str | None = None) -> str:
    card_class = "card intro-card" if num is None else "card spec-card"
    bar = metric_emphasis(metric) if metric else ""
    return f'<section class="{card_class}">{section_head(num, title)}<div class="gray-panel spec-text">{bar}{body}</div></section>'


# Shared helper: html2canvas does NOT support object-fit or mix-blend-mode, so a
# replaced .hero-overlay image would download deformed (stretched) and un-blended
# (wrong colour). Before capture we pre-composite the red + overlay image into a
# single flat, hero-sized PNG (cover-fit + the blend baked in via canvas
# globalCompositeOperation), swap it in (no blend, fill = no distortion), and
# restore afterwards — so the download matches the on-screen preview.
HERO_FLATTEN_JS = (
    "function flattenHeroOverlay(doc){return new Promise(function(resolve){try{"
    "var ov=doc.querySelector('.hero-overlay');"
    "if(!ov||!ov.naturalWidth||ov.naturalWidth<=2){return resolve(null);}"
    "var hero=ov.closest('.hero')||ov.parentNode;var r=hero.getBoundingClientRect();"
    "var W=Math.max(1,Math.round(r.width)),H=Math.max(1,Math.round(r.height));"
    "var win=doc.defaultView||window,cs=win.getComputedStyle(ov);"
    "var blend=(cs.mixBlendMode&&cs.mixBlendMode!=='normal')?cs.mixBlendMode:'overlay';"
    "var bg=win.getComputedStyle(hero).backgroundColor||'#FF2B22';"
    "var cnv=doc.createElement('canvas');cnv.width=W;cnv.height=H;var ctx=cnv.getContext('2d');"
    "ctx.fillStyle=bg;ctx.fillRect(0,0,W,H);"
    "var iw=ov.naturalWidth,ih=ov.naturalHeight,sc=Math.max(W/iw,H/ih),dw=iw*sc,dh=ih*sc;"
    "ctx.globalCompositeOperation=blend;ctx.drawImage(ov,(W-dw)/2,(H-dh)/2,dw,dh);"
    "ctx.globalCompositeOperation='source-over';var flat=cnv.toDataURL('image/png');"
    "var saved={src:ov.getAttribute('src'),mb:ov.style.mixBlendMode,of:ov.style.objectFit};"
    "var fired=false;var done=function(){if(fired)return;fired=true;ov.removeEventListener('load',done);"
    "resolve(function(){ov.setAttribute('src',saved.src);ov.style.mixBlendMode=saved.mb;ov.style.objectFit=saved.of;});};"
    "ov.addEventListener('load',done);ov.style.mixBlendMode='normal';ov.style.objectFit='fill';"
    "ov.setAttribute('src',flat);setTimeout(done,800);}catch(e){resolve(null);}});}"
)

# html2canvas has no <video> handling: a video only renders if the browser already
# decoded a frame. During capture, swap each <video> for an <img> of its poster
# (or decoded frame) so the video's cover always appears, then restore.
VIDEO_SWAP_JS = (
    "function swapVideosForPosters(doc){var swaps=[];"
    "[].slice.call(doc.querySelectorAll('video')).forEach(function(v){"
    "var src=v.getAttribute('poster')||'';"
    "if(!src&&v.videoWidth){try{var c=doc.createElement('canvas');c.width=v.videoWidth;c.height=v.videoHeight;"
    "c.getContext('2d').drawImage(v,0,0,c.width,c.height);src=c.toDataURL('image/jpeg',0.85);}catch(e){}}"
    "if(!src)return;var img=doc.createElement('img');img.src=src;img.className=v.className;"
    "img.style.cssText=v.style.cssText;img.style.display='block';img.style.width=v.style.width||'100%';img.style.height='auto';"
    "v.parentNode.insertBefore(img,v);var pd=v.style.display;v.style.display='none';swaps.push({v:v,img:img,pd:pd});});"
    "return function(){swaps.forEach(function(s){if(s.img.parentNode)s.img.parentNode.removeChild(s.img);s.v.style.display=s.pd;});};}"
)


def download_runtime() -> str:
    """A floating "下载整页图片" button that rasterises the whole poster to one
    PNG via an embedded html2canvas, so the page stays self-contained/offline.
    Scale is clamped so the canvas height stays under the browser limit."""
    if not DEFAULT_H2C.exists():
        return ""
    lib = DEFAULT_H2C.read_text(encoding="utf-8")
    return (
        '<button id="dl-page-btn" class="dl-page-btn" data-html2canvas-ignore>下载整页图片</button>\n'
        f"<script>{lib}</script>\n"
        "<script>(function(){var b=document.getElementById('dl-page-btn');"
        "if(!b||!window.html2canvas)return;"
        f"{HERO_FLATTEN_JS}{VIDEO_SWAP_JS}"
        "b.addEventListener('click',function(){var p=document.querySelector('.poster');if(!p)return;"
        "var t=b.textContent;b.textContent='生成中…';b.disabled=true;var restore=null,vrestore=null;"
        "flattenHeroOverlay(document).then(function(rf){restore=rf;vrestore=swapVideosForPosters(document);"
        "var h=p.scrollHeight,s=Math.min(2,Math.max(0.3,32000/h));"
        "return html2canvas(p,{scale:s,backgroundColor:'#dcedff',useCORS:true,logging:false,"
        "windowWidth:p.scrollWidth,windowHeight:h});}).then(function(c){if(vrestore)vrestore();if(restore)restore();c.toBlob(function(bl){"
        "var a=document.createElement('a');a.href=URL.createObjectURL(bl);"
        "a.download=(document.title||'page')+'.png';document.body.appendChild(a);a.click();a.remove();"
        "setTimeout(function(){URL.revokeObjectURL(a.href);},1500);b.textContent=t;b.disabled=false;},'image/png');"
        "}).catch(function(e){if(vrestore)vrestore();if(restore)restore();console.error(e);b.textContent='下载失败，重试';b.disabled=false;});});})();</script>"
    )


def editor_runtime() -> str:
    """A floating 「编辑」 button that opens the bundled visual HTML editor in a NEW
    window, pre-loaded with the current page. The whole editor is embedded as inert
    base64 so the page stays self-contained / offline (no second file needed)."""
    if not DEFAULT_EDITOR.exists():
        return ""
    b64 = base64.b64encode(DEFAULT_EDITOR.read_bytes()).decode("ascii")
    launcher = (
        "(function(){var b=document.getElementById('edit-page-btn');"
        "var srcEl=document.getElementById('editor-src-b64');"
        "if(!b||!srcEl)return;"
        "b.addEventListener('click',function(){"
        # Clean snapshot of the current poster: drop the floating controls + scripts.
        "var root=document.documentElement.cloneNode(true);"
        "root.querySelectorAll('[data-html2canvas-ignore],script').forEach(function(n){n.remove();});"
        "var poster='<!doctype html>\\n'+root.outerHTML;"
        # Decode the embedded editor (UTF-8 safe), then run it in a new window with
        # the current page handed off via window.__PRELOAD_HTML__.
        "var bin=atob(srcEl.textContent.trim());var bytes=new Uint8Array(bin.length);"
        "for(var i=0;i<bin.length;i++){bytes[i]=bin.charCodeAt(i);}"
        "var editorHtml=new TextDecoder('utf-8').decode(bytes);"
        "var name=(document.title||'页面')+'.html';"
        "var boot='<script>window.__PRELOAD_NAME__='+JSON.stringify(name)+';window.__PRELOAD_HTML__='+JSON.stringify(poster)+';<\\/script>';"
        # Inject boot inside <head> so the doctype stays first (no quirks mode).
        "var out=editorHtml.replace(/<head([^>]*)>/i,function(m){return m+boot;});"
        "if(out===editorHtml){out=boot+editorHtml;}"
        "var w=window.open('','_blank');"
        "if(!w){alert('请允许弹出窗口后重试');return;}"
        "w.document.open();w.document.write(out);w.document.close();"
        "});})();"
    )
    return (
        '<button id="edit-page-btn" class="edit-page-btn" data-html2canvas-ignore>编辑</button>\n'
        f'<script type="application/octet-stream" id="editor-src-b64" data-html2canvas-ignore>{b64}</script>\n'
        f"<script>{launcher}</script>"
    )


def render_html(docx_path: Path, design_path: Path, font_path: Path | None, updated_label: str, editable: bool) -> str:
    doc = Document(docx_path)
    blobs = image_target_to_blob(doc)
    title, sections = split_sections(iter_blocks(doc))
    css = extract_css(design_path, font_path)
    cards: list[str] = []
    chapter = 1
    for section_title, section_blocks, section_metric in sections:
        is_intro = section_title == "整体规范综述" and not cards
        half_images = "图文详情" in section_title
        video_section = "主图视频" in section_title
        body = render_section_blocks(
            section_blocks, doc, blobs, is_intro=is_intro, half_images=half_images,
            video_section=video_section,
        )
        if is_intro:
            cards.append(render_card(section_title, body, None))
        else:
            cards.append(render_card(section_title, body, chapter, section_metric))
            chapter += 1
    editable_runtime = EDITABLE_RUNTIME if editable else ""
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{esc(title)}</title>
  <style>{css}</style>
</head>
<body>
<main class="poster auto-doc">
  <section class="hero">
    <img class="hero-overlay" src="{HERO_OVERLAY_SRC}" alt="头图背景图（在编辑器中双击此处可替换为叠加在红底上的图片）">
    <h1>{hero_title_html(title)}</h1>
    <p class="updated">{esc(updated_label)}</p>
    <div class="hero-mark">OPERATION<br>STANDARDS</div>
    <div class="hero-rule">{HERO_RULE_ARROW_SVG}</div>
  </section>
  {''.join(cards)}
</main>
{download_runtime()}
{editor_runtime()}
{editable_runtime}
</body>
</html>
"""


def docx_inputs(input_path: Path) -> list[Path]:
    if input_path.is_file():
        return [input_path]
    return sorted(path for path in input_path.rglob("*.docx") if not path.name.startswith("~$"))


def generate_one(docx_path: Path, output_dir: Path, design_path: Path, font_path: Path | None, updated_value: str | None, editable: bool) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    out_html = output_dir / f"{slugify(docx_path.stem)}-output.html"
    report_path = output_dir / f"{slugify(docx_path.stem)}-report.json"
    updated_label = normalize_update_label(updated_value, docx_path)
    html_text = render_html(docx_path, design_path, font_path, updated_label, editable)
    out_html.write_text(html_text, encoding="utf-8")
    report = validate(docx_path, out_html)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return {
        "source": str(docx_path),
        "html": str(out_html),
        "report": str(report_path),
        "passed": report["passed"],
        "warnings": report["warnings"],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Batch-generate single-file HTML from DOCX files.")
    parser.add_argument("input", type=Path, help="A .docx file or a folder containing .docx files.")
    parser.add_argument("output_dir", type=Path, help="Folder for generated HTML and reports.")
    parser.add_argument("--design", type=Path, default=DEFAULT_DESIGN)
    parser.add_argument("--font", type=Path, default=DEFAULT_FONT)
    parser.add_argument("--updated", default=None, help='Hero update label. Examples: "2026.06" or "更新日期 2026年06月".')
    parser.add_argument("--editable", action="store_true", help="Add an optional in-page text editing toolbar and download button.")
    parser.add_argument("--strict", action="store_true", help="Exit non-zero if any generated report has warnings.")
    args = parser.parse_args()

    inputs = docx_inputs(args.input)
    if not inputs:
        raise SystemExit(f"No .docx files found: {args.input}")

    font_path = args.font if args.font and args.font.exists() else None
    summary = [generate_one(path, args.output_dir, args.design, font_path, args.updated, args.editable) for path in inputs]
    summary_path = args.output_dir / "batch-summary.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"count": len(summary), "summary": str(summary_path), "items": summary}, ensure_ascii=False, indent=2))
    return 1 if args.strict and any(not item["passed"] for item in summary) else 0


if __name__ == "__main__":
    raise SystemExit(main())
