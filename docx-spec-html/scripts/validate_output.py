#!/usr/bin/env python3
from __future__ import annotations

import argparse
import html.parser
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document


def clean_text(value: str) -> str:
    value = re.sub(r"[\u200b-\u200f\ufeff]", "", value or "")
    return re.sub(r"\s+", " ", value).strip()


class VisibleTextParser(html.parser.HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"style", "script", "noscript"}:
            self.skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"style", "script", "noscript"} and self.skip_depth:
            self.skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    @property
    def text(self) -> str:
        return clean_text("".join(self.parts))


def docx_texts(docx_path: Path) -> list[str]:
    doc = Document(docx_path)
    values: list[str] = []
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            values.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = clean_text(paragraph.text)
                    if text:
                        values.append(text)
    return values


def docx_image_count(docx_path: Path) -> int:
    doc = Document(docx_path)
    return len(doc.element.body.xpath(".//a:blip"))


def docx_table_count(docx_path: Path) -> int:
    return len(Document(docx_path).tables)


def html_visible_text(html_text: str) -> str:
    parser = VisibleTextParser()
    parser.feed(html_text)
    return parser.text


def strip_tags(value: str) -> str:
    return clean_text(re.sub(r"<[^>]+>", "", value))


def count_class(html_text: str, class_name: str) -> int:
    pattern = rf'<[^>]+\bclass="[^"]*\b{re.escape(class_name)}\b[^"]*"'
    return len(re.findall(pattern, html_text, flags=re.I))


def structural_checks(html_text: str, visible: str, expected_images: int) -> tuple[dict[str, bool], list[str]]:
    h2_texts = [
        strip_tags(match)
        for match in re.findall(r"<h2\b[^>]*>(.*?)</h2>", html_text, flags=re.I | re.S)
    ]
    section_head_count = count_class(html_text, "section-head")
    spec_head_count = len(
        re.findall(
            r'<div\b[^>]*class="(?=[^"]*\bsection-head\b)(?=[^"]*\bspec-head\b)[^"]*"',
            html_text,
            flags=re.I,
        )
    )
    en_label_count = count_class(html_text, "en-label")
    label_line_count = count_class(html_text, "label-line")
    label_text_count = count_class(html_text, "label-text")
    text_block_count = count_class(html_text, "text-block")
    caption_card_count = count_class(html_text, "caption-image-card")
    image_holder_count = count_class(html_text, "image-holder")
    image_frame_count = count_class(html_text, "image-frame")
    screens_grid_count = count_class(html_text, "screens-grid")
    detail_screen_grid_count = count_class(html_text, "detail-screen-grid")
    screen_count = len(re.findall(r"第\s*\d+\s*屏|【第\s*\d+\s*屏】", visible))
    formula_caption_like = bool(re.search(r"\[[^\]]+\]\s*\+", visible))

    checks = {
        "section_title_single_spaced_braces": bool(h2_texts)
        and all(re.fullmatch(r"\{ [^{}]+ \}", text) for text in h2_texts),
        "section_heads_include_spec_head_and_label": section_head_count > 0
        and section_head_count == spec_head_count == en_label_count,
        "label_lines_wrap_label_text": label_line_count == 0 or label_line_count == label_text_count,
        "white_text_blocks_used_for_label_groups": label_line_count < 6 or text_block_count > 0,
        "image_cards_use_image_holders": expected_images == 0 or image_holder_count >= expected_images,
        "caption_images_keep_text_below_images": not formula_caption_like or caption_card_count > 0,
        "many_screen_examples_use_detail_grid": screen_count < 5 or detail_screen_grid_count > 0,
        "detail_grid_not_four_column_screens_grid": screen_count < 5
        or screens_grid_count == 0,
    }

    messages = {
        "section_title_single_spaced_braces": "Section h2 titles must be exactly `{ 标题 }`; double braces or missing inner spaces are not allowed.",
        "section_heads_include_spec_head_and_label": "Every section-head, including overview cards, must use spec-head and include the right INTRODUCTION label.",
        "label_lines_wrap_label_text": "Every label-line must contain label-text so the pink highlight stays behind text only.",
        "white_text_blocks_used_for_label_groups": "Grey-panel label groups must be wrapped in white text-block modules instead of bare labels/lists.",
        "image_cards_use_image_holders": "Every DOCX image should be placed inside an image-holder so it keeps proportion and spacing.",
        "caption_images_keep_text_below_images": "Formula or caption text after an image must use caption-image-card, with image above and caption below.",
        "many_screen_examples_use_detail_grid": "Five or more screen/detail examples must use detail-screen-grid instead of a vertical list or cramped grid.",
        "detail_grid_not_four_column_screens_grid": "Screen/detail examples must not use the four-column screens-grid pattern; use the two-column detail-screen-grid.",
    }
    warnings = [message for key, message in messages.items() if not checks[key]]
    return checks, warnings


def validate(docx_path: Path, html_path: Path) -> dict[str, Any]:
    html_text = html_path.read_text(encoding="utf-8")
    visible = html_visible_text(html_text)
    source_texts = docx_texts(docx_path)
    expected_counts = Counter(source_texts)
    actual_counts = Counter()
    missing: list[str] = []
    underrepresented: list[dict[str, Any]] = []

    for text, expected in expected_counts.items():
        actual = visible.count(text)
        actual_counts[text] = actual
        if actual == 0:
            missing.append(text)
        elif actual < expected:
            underrepresented.append({"text": text, "expected": expected, "actual": actual})

    expected_images = docx_image_count(docx_path)
    actual_images = len(re.findall(r"<img\b", html_text, flags=re.I))
    expected_tables = docx_table_count(docx_path)
    actual_table_like = (
        len(re.findall(r"<table\b", html_text, flags=re.I))
        + html_text.count("word-table-spec")
        + html_text.count("doc-table")
    )

    css_checks = {
        "embedded_style": "<style>" in html_text,
        "embedded_font_face": "@font-face" in html_text,
        "hero_font": "JINGDONGLangZhengTi1-Bold" in html_text,
        "hero_mark_present": "OPERATION" in html_text and "STANDARDS" in html_text and "hero-mark" in html_text,
        "chinese_update_label": "更新日期" in html_text,
        "introduction_labels": "<strong>INTRODUCTION</strong>" in html_text,
        "panel_spacing_rule": ".gray-panel > * + *" in html_text,
        "single_file_images": "data:image/" in html_text if expected_images else True,
        "old_content_labels_absent": not any(
            value in html_text
            for value in [
                "MAIN IMAGE",
                "MAIN VIDEO",
                "LONG TITLE",
                "SHORT TITLE",
                "SELLING POINT",
                "PARAMETER",
                "DETAIL PAGE",
            ]
        ),
    }
    structure_checks, structure_warnings = structural_checks(html_text, visible, expected_images)

    warnings: list[str] = []
    if missing:
        warnings.append("Some DOCX text fragments are missing from visible HTML text.")
    if underrepresented:
        warnings.append("Some repeated DOCX text fragments appear fewer times in HTML.")
    if actual_images < expected_images:
        warnings.append("HTML image count is lower than DOCX image occurrence count.")
    if expected_tables and actual_table_like < expected_tables:
        warnings.append("DOCX contains tables; HTML may not preserve all table-like structures.")
    if not all(css_checks.values()):
        warnings.append("One or more required CSS or layout invariants are missing.")
    warnings.extend(structure_warnings)

    return {
        "source": str(docx_path),
        "html": str(html_path),
        "source_text_count": len(source_texts),
        "unique_source_text_count": len(expected_counts),
        "missing_text_count": len(missing),
        "missing_texts": missing[:50],
        "underrepresented_text_count": len(underrepresented),
        "underrepresented_texts": underrepresented[:50],
        "expected_image_count": expected_images,
        "actual_image_count": actual_images,
        "expected_table_count": expected_tables,
        "actual_table_like_count": actual_table_like,
        "css_checks": css_checks,
        "structure_checks": structure_checks,
        "passed": not warnings,
        "warnings": warnings,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate DOCX-to-MPDN50EU HTML output.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("html", type=Path)
    parser.add_argument("--report", type=Path)
    parser.add_argument("--strict", action="store_true", help="Exit non-zero when validation warnings exist.")
    args = parser.parse_args()

    report = validate(args.docx, args.html)
    text = json.dumps(report, ensure_ascii=False, indent=2)
    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(text + "\n", encoding="utf-8")
    print(text)
    return 1 if args.strict and not report["passed"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
